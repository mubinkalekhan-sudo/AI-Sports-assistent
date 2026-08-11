"""
SportsAI - Premium AI Sports Assistant Chatbot
--------------------------------------------------------
A Streamlit app that uses Anthropic Claude and Google Gemini APIs to power an
interactive, visually stunning sports training, rules, quiz, and QA assistant.

Features:
- Dual LLM support (Anthropic Claude + Google Gemini)
- Interactive Quiz game
- Weekly Training Plan generator & Word Document (.docx) exporter
- RAG document search (PDF/DOCX) using FAISS & Sentence Transformers
"""

import os
import streamlit as st
import io
from dotenv import load_dotenv
from streamlit_mic_recorder import speech_to_text

# ------------------------------------------------------------------
# Conditional imports for LLM, document parsing, and RAG libraries
# ------------------------------------------------------------------
try:
    from anthropic import Anthropic
    HAS_ANTHROPIC = True
except ImportError:
    HAS_ANTHROPIC = False

try:
    import google.generativeai as genai
    HAS_GEMINI = True
except ImportError:
    HAS_GEMINI = False

try:
    from sentence_transformers import SentenceTransformer
    import faiss
    import numpy as np
    HAS_RAG = True
except ImportError:
    HAS_RAG = False

try:
    import pypdf
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

import time
import random

# Load environment variables robustly (detecting standard and raw keys)
def load_keys_robust():
    load_dotenv(override=True)
    gemini_key = os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY", "")
    anthropic_key = os.getenv("ANTHROPIC_API_KEY", "")
    
    # If keys are empty, try manual parsing of raw keys in .env
    if os.path.exists(".env"):
        try:
            with open(".env", "r", encoding="utf-8") as f:
                lines = [line.strip() for line in f.read().splitlines() if line.strip()]
                if lines:
                    # Case 1: Just a raw key on the first line (no equals sign)
                    if len(lines) == 1 and "=" not in lines[0]:
                        raw_key = lines[0]
                        if raw_key.startswith("sk-ant-"):
                            if not anthropic_key:
                                anthropic_key = raw_key
                        else:
                            if not gemini_key:
                                gemini_key = raw_key
                            if not anthropic_key:
                                anthropic_key = raw_key
                    else:
                        # Case 2: Standard or mixed lines
                        for line in lines:
                            if "=" in line:
                                parts = line.split("=", 1)
                                k = parts[0].strip()
                                v = parts[1].strip().strip('"').strip("'")
                                if k == "GEMINI_API_KEY" or k == "GOOGLE_API_KEY":
                                    if not gemini_key:
                                        gemini_key = v
                                elif k == "ANTHROPIC_API_KEY":
                                    if not anthropic_key:
                                        anthropic_key = v
                            else:
                                # Fallback raw key detection
                                val = line.strip()
                                if val.startswith("sk-ant-"):
                                    if not anthropic_key:
                                        anthropic_key = val
                                elif len(val) > 20 and not val.startswith("#"):
                                    if not gemini_key:
                                        gemini_key = val
        except Exception:
            pass
    return gemini_key, anthropic_key

ENV_GEMINI_KEY, ENV_ANTHROPIC_KEY = load_keys_robust()
DEFAULT_MODEL_GEMINI = os.getenv("GEMINI_MODEL", "gemini-3.5-flash")
DEFAULT_MODEL_CLAUDE = os.getenv("MODEL_NAME", "claude-3-5-sonnet-latest")

# Initialize theme session state
if "current_theme" not in st.session_state:
    st.session_state.current_theme = "Rose Quartz Light"
if "ai_chat_color" not in st.session_state:
    st.session_state.ai_chat_color = "Championship Gold"

# Initialize chat sessions
if "chat_sessions" not in st.session_state:
    st.session_state.chat_sessions = {"default": {"name": "Default Chat", "messages": []}}
if "current_session_id" not in st.session_state:
    st.session_state.current_session_id = "default"

# Initialize Career & Guidance session states
if "career_goals" not in st.session_state:
    st.session_state.career_goals = []
if "career_achievements" not in st.session_state:
    st.session_state.career_achievements = []
if "career_chatbot_messages" not in st.session_state:
    st.session_state.career_chatbot_messages = []
if "career_roadmap" not in st.session_state:
    st.session_state.career_roadmap = None
if "skill_assessment_results" not in st.session_state:
    st.session_state.skill_assessment_results = None

# Sync messages with active session
st.session_state.messages = st.session_state.chat_sessions[st.session_state.current_session_id]["messages"]

SYSTEM_PROMPT = """You are SportsAI Pro, an advanced AI-powered Sports Assistant designed to provide accurate, engaging, and personalized support for sports fans, athletes, students, coaches, referees, and fitness enthusiasts.

Your mission is to become the user's personal sports coach, analyst, trainer, statistician, commentator, fitness guide, and sports encyclopedia.

Always provide clear, accurate, friendly, and well-structured responses.

---

# 🎯 Core Features

## ⚽ 1. Sports Knowledge
Answer questions about: Cricket, Football, Basketball, Tennis, Badminton, Volleyball, Hockey, Kabaddi, Baseball, Rugby, Athletics, Swimming, Formula 1, Golf, Boxing, MMA, Chess, Esports, Olympics, Paralympics.
Explain: Rules, History, Equipment, Positions, Scoring, Fouls, Strategies, Techniques.

## 📅 2. Match Schedule
Provide: Today's matches, Upcoming matches, Fixtures, Match timings, Stadium, Tournament schedule, League schedule.

## 📺 3. Live Match Assistant
When real-time information is available: Provide Live score, Overs, Goals, Cards, Possession, Corners, Match statistics, Current standings. If live data is unavailable, clearly say so instead of guessing.

## 📊 4. Player Statistics
Show: Matches played, Runs, Goals, Assists, Wickets, Strike Rate, Batting Average, Bowling Average, Rankings, Career Records, Awards, Milestones.

## 🏆 5. Team Analysis
Provide: Team overview, Squad, Captain, Coach, Playing XI, Key players, Strengths, Weaknesses, Recent form, Tactical analysis.

## 📈 6. Match Prediction
Predict using: Team form, Home advantage, Head-to-head record, Injuries, Weather, Pitch conditions. Always explain that predictions are probabilistic, not guaranteed.

## 🌦️ 7. Weather & Pitch Report
Provide: Match weather, Temperature, Rain chances, Pitch condition, Wind, Humidity. Explain how these conditions may affect play.

## 📰 8. Sports News
Summarize: Breaking sports news, Transfer news, Injury updates, Tournament news, Team announcements, Award news. Indicate when information may not be current.

## 🏅 9. Tournament Guide
Support: IPL, ICC World Cup, Champions Trophy, FIFA World Cup, UEFA Champions League, EPL, La Liga, Serie A, Bundesliga, NBA, Wimbledon, Olympics, Asian Games, Commonwealth Games, Pro Kabaddi, ISL.

## 🏃 10. Personal Sports Coach
Create personalized: Daily practice plans, Weekly schedules, Monthly plans, Warm-up routines, Cool-down routines, Stretching, Recovery exercises.

## 💪 11. Fitness Trainer
Generate: Weight loss workouts, Muscle building plans, Cardio plans, Strength training, Speed drills, Agility drills, Endurance training.

## 🍎 12. Sports Nutrition
Recommend: Healthy meals, Hydration, Protein intake, Pre-workout meals, Post-workout meals, Recovery nutrition. Do not provide medical diagnoses.

## 🎯 13. Skill Improvement
Help improve: Batting, Bowling, Goalkeeping, Shooting, Passing, Dribbling, Serving, Sprinting, Jumping, Reflexes.

## 🎥 14. Video & Drill Suggestions
Recommend the types of training videos or drills users should look for to improve specific skills.

## 🧠 15. Sports Quiz Mode
Generate: Easy quizzes, Medium quizzes, Hard quizzes, MCQs, True/False, Fill in the blanks.

## 🎮 16. Fun Mode
Create: Guess the player, Guess the logo, Guess the stadium, Sports riddles, Trivia, Fun facts.

## 📖 17. Sports Dictionary
Explain sports terminology in simple language (e.g. Offside, LBW, Hat-trick, Free Kick, Yorker, Power Play).

## 🏃‍♂️ 18. Injury Prevention
Provide general advice on: Warm-up, Stretching, Recovery, Rest, Safe training habits. Recommend consulting a healthcare professional for injuries.

## 📅 19. Training Planner
Generate: Daily planner, Weekly planner, Monthly planner, Tournament preparation schedule.

## 🏅 20. Athlete Profile
Create athlete summaries including: Biography, Achievements, Career timeline, Playing style, Major records.

## 📊 21. Comparison Mode
Compare: Players, Teams, Coaches, Stadiums, Leagues, Tournaments. Present comparisons in a clear table whenever appropriate.

## 🌍 22. Multi-Language Support
Respond in the user's preferred language.

## 🎙️ 23. Commentary Mode
Generate: Ball-by-ball style commentary, Match summaries, Post-match analysis, Highlights. Make it clear when commentary is simulated rather than live.

## 🤖 24. AI Coach Mode
Ask users: Age, Height, Weight, Sport, Experience, Goal. Then generate a personalized improvement plan.

## 📚 25. Sports Learning Mode
Teach beginners: Rules, Skills, Techniques, Strategies. Use simple language and step-by-step explanations.

## ⏳ 26. Sports History Explorer
Provide extensive, engaging, and accurate historical sports information. Highlight key historical milestones, rules evolution (e.g. how the offside rule or LBW rule evolved), famous matches, legendary rivalries, and biographical details of iconic athletes. When describing historical matches, use a storytelling tone, details on tactics, and key stats.

## 🧭 27. Sports Career Advisor
Provide career roadmaps, guidance, licensing paths, and advice for sports-related professions. This includes professional athlete paths, sports coaching license requirements (such as UEFA/AFC for football, BCCI/ICC for cricket), sports analytics, sports medicine/physiotherapy, sports journalism, and sports management. Give step-by-step development guidance.

---

# 🎨 Response Style
Always:
- Use clear headings.
- Use bullet points and tables where helpful.
- Explain concepts in simple language.
- Ask follow-up questions when needed.
- Never fabricate live scores or current events.
- Clearly distinguish facts from opinions or predictions.
- Keep responses motivating and sportsmanlike.

---

# Closing Behavior
End every response with the exact sentence:
"Would you like player statistics, match analysis, training plans, fitness advice, sports quizzes, tournament information, sports history, or career guidance? I'm ready to help!"
"""

# ------------------------------------------------------------------
# Page config & Custom Premium Dark-Neon Glassmorphic CSS Styling
# ------------------------------------------------------------------
st.set_page_config(
    page_title="SportsAI - Premium Sports Assistant",
    page_icon="🏆",
    layout="wide",
)

# ------------------------------------------------------------------
# Page config & Custom Premium CSS Styling (Rose Quartz vs. Cyberpunk Dark)
# ------------------------------------------------------------------
ROSE_QUARTZ_CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;600;700&display=swap');

html, body, [class*="css"], .stApp {
    font-family: 'Outfit', sans-serif;
    background-color: #fff0f5 !important;
    color: #000000 !important;
}

section[data-testid="stSidebar"] {
    background: #ffe4e1 !important;
    border-right: 1px solid rgba(255, 105, 180, 0.3) !important;
    backdrop-filter: blur(16px) !important;
    -webkit-backdrop-filter: blur(16px) !important;
}

section[data-testid="stSidebar"] h1, 
section[data-testid="stSidebar"] h2, 
section[data-testid="stSidebar"] h3,
section[data-testid="stSidebar"] label {
    color: #ff1493 !important;
    font-weight: 600;
}

section[data-testid="stSidebar"] p,
section[data-testid="stSidebar"] span,
section[data-testid="stSidebar"] div {
    color: #000000 !important;
}

.main-title {
    font-size: 3.2rem;
    font-weight: 800;
    background: linear-gradient(135deg, #ff1493 0%, #ff69b4 100%);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    margin-bottom: 0.2rem;
    margin-top: -1rem;
    letter-spacing: -1px;
}

.subtitle {
    font-size: 1.2rem;
    color: #000000;
    margin-bottom: 2rem;
}

.welcome-card {
    background: #ffffff !important;
    border: 1px solid rgba(255, 105, 180, 0.25);
    border-radius: 16px;
    padding: 2.2rem;
    margin-bottom: 2rem;
    box-shadow: 0 10px 30px rgba(255, 105, 180, 0.08);
}

.welcome-header {
    font-size: 1.7rem;
    font-weight: 700;
    color: #ff1493;
    margin-bottom: 1rem;
}

.welcome-item {
    display: flex;
    align-items: center;
    margin-bottom: 0.6rem;
    font-size: 1.05rem;
    color: #000000;
}

.welcome-icon {
    color: #ff1493;
    margin-right: 0.8rem;
    font-weight: bold;
}

.premium-card {
    background: #ffffff !important;
    border: 1px solid rgba(255, 105, 180, 0.15);
    border-radius: 12px;
    padding: 1.5rem;
    margin-bottom: 1.2rem;
    box-shadow: 0 4px 12px rgba(255, 105, 180, 0.05);
    transition: all 0.3s ease;
    color: #000000 !important;
}

.premium-card:hover {
    transform: translateY(-2px);
    border-color: rgba(255, 105, 180, 0.5);
    box-shadow: 0 8px 24px rgba(255, 105, 180, 0.15);
}

.stButton > button {
    background: linear-gradient(135deg, #ff69b4 0%, #ff1493 100%) !important;
    color: #ffffff !important;
    border: none !important;
    border-radius: 8px !important;
    padding: 0.5rem 1.8rem !important;
    font-weight: 600 !important;
    transition: all 0.3s ease !important;
    box-shadow: 0 4px 14px rgba(255, 105, 180, 0.3) !important;
}

.stButton > button:hover {
    transform: translateY(-2px) !important;
    box-shadow: 0 6px 20px rgba(255, 105, 180, 0.5) !important;
    background: linear-gradient(135deg, #ff1493 0%, #ff69b4 100%) !important;
    color: #ffffff !important;
}

.stChatMessage {
    border-radius: 16px !important;
    margin-bottom: 1rem !important;
    padding: 1.2rem !important;
    transition: all 0.3s ease !important;
}

.stChatMessage[data-testid="stChatMessageUser"] {
    background: linear-gradient(135deg, rgba(255, 182, 193, 0.3) 0%, rgba(255, 240, 245, 0.3) 100%) !important;
    border: 1px solid rgba(255, 105, 180, 0.25) !important;
    box-shadow: 0 4px 15px rgba(255, 105, 180, 0.06) !important;
    color: #000000 !important;
}

/* Custom Chat Input & General Text Box Styling */
div[data-testid="stChatInput"] {
    background-color: transparent !important;
    padding: 0 !important;
}
div[data-testid="stChatInput"] textarea {
    background-color: #ffffff !important;
    color: #000000 !important;
    border: 2px solid #ff69b4 !important;
    border-radius: 14px !important;
    box-shadow: 0 4px 15px rgba(255, 105, 180, 0.15) !important;
}
div[data-testid="stChatInput"] textarea:focus {
    border-color: #ff1493 !important;
    box-shadow: 0 4px 20px rgba(255, 105, 180, 0.3) !important;
}
.stTextInput input {
    background-color: #ffffff !important;
    color: #000000 !important;
    border: 1px solid rgba(255, 105, 180, 0.3) !important;
    border-radius: 8px !important;
    box-shadow: 0 2px 8px rgba(255, 105, 180, 0.05) !important;
}
.stTextInput input:focus {
    border-color: #ff1493 !important;
    box-shadow: 0 2px 12px rgba(255, 105, 180, 0.15) !important;
}

/* Fix input element labels to remain dark */
label[data-testid="stWidgetLabel"], .stSelectbox label, .stTextInput label, .stSlider label, .stCheckbox label {
    color: #000000 !important;
}
</style>
"""

CYBERPUNK_DARK_CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;600;700&display=swap');

html, body, [class*="css"], .stApp {
    font-family: 'Outfit', sans-serif;
    background-color: #0a0e17 !important;
    color: #f1f5f9 !important;
}

section[data-testid="stSidebar"] {
    background: #0f172a !important;
    border-right: 1px solid rgba(6, 182, 212, 0.3) !important;
    backdrop-filter: blur(16px) !important;
    -webkit-backdrop-filter: blur(16px) !important;
}

section[data-testid="stSidebar"] h1, 
section[data-testid="stSidebar"] h2, 
section[data-testid="stSidebar"] h3,
section[data-testid="stSidebar"] label {
    color: #06b6d4 !important;
    font-weight: 600;
}

section[data-testid="stSidebar"] p,
section[data-testid="stSidebar"] span,
section[data-testid="stSidebar"] div {
    color: #f1f5f9 !important;
}

.main-title {
    font-size: 3.2rem;
    font-weight: 800;
    background: linear-gradient(135deg, #06b6d4 0%, #d946ef 100%);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    margin-bottom: 0.2rem;
    margin-top: -1rem;
    letter-spacing: -1px;
}

.subtitle {
    font-size: 1.2rem;
    color: #94a3b8;
    margin-bottom: 2rem;
}

.welcome-card {
    background: #1e293b !important;
    border: 1px solid rgba(6, 182, 212, 0.25);
    border-radius: 16px;
    padding: 2.2rem;
    margin-bottom: 2rem;
    box-shadow: 0 10px 30px rgba(6, 182, 212, 0.08);
}

.welcome-header {
    font-size: 1.7rem;
    font-weight: 700;
    color: #06b6d4;
    margin-bottom: 1rem;
}

.welcome-item {
    display: flex;
    align-items: center;
    margin-bottom: 0.6rem;
    font-size: 1.05rem;
    color: #cbd5e1;
}

.welcome-icon {
    color: #06b6d4;
    margin-right: 0.8rem;
    font-weight: bold;
}

.premium-card {
    background: #1e293b !important;
    border: 1px solid rgba(6, 182, 212, 0.15);
    border-radius: 12px;
    padding: 1.5rem;
    margin-bottom: 1.2rem;
    box-shadow: 0 4px 12px rgba(6, 182, 212, 0.05);
    transition: all 0.3s ease;
    color: #f1f5f9 !important;
}

.premium-card:hover {
    transform: translateY(-2px);
    border-color: rgba(6, 182, 212, 0.5);
    box-shadow: 0 8px 24px rgba(6, 182, 212, 0.15);
}

.stButton > button {
    background: linear-gradient(135deg, #06b6d4 0%, #d946ef 100%) !important;
    color: #ffffff !important;
    border: none !important;
    border-radius: 8px !important;
    padding: 0.5rem 1.8rem !important;
    font-weight: 600 !important;
    transition: all 0.3s ease !important;
    box-shadow: 0 4px 14px rgba(6, 182, 212, 0.3) !important;
}

.stButton > button:hover {
    transform: translateY(-2px) !important;
    box-shadow: 0 6px 20px rgba(6, 182, 212, 0.5) !important;
    background: linear-gradient(135deg, #d946ef 0%, #06b6d4 100%) !important;
    color: #ffffff !important;
}

.stChatMessage {
    border-radius: 16px !important;
    margin-bottom: 1rem !important;
    padding: 1.2rem !important;
    transition: all 0.3s ease !important;
}

.stChatMessage[data-testid="stChatMessageUser"] {
    background: linear-gradient(135deg, rgba(30, 41, 59, 0.8) 0%, rgba(15, 23, 42, 0.8) 100%) !important;
    border: 1px solid rgba(6, 182, 212, 0.25) !important;
    box-shadow: 0 4px 15px rgba(6, 182, 212, 0.06) !important;
    color: #f1f5f9 !important;
}

/* Custom Chat Input & General Text Box Styling */
div[data-testid="stChatInput"] {
    background-color: transparent !important;
    padding: 0 !important;
}
div[data-testid="stChatInput"] textarea {
    background-color: #0f172a !important;
    color: #f1f5f9 !important;
    border: 2px solid #06b6d4 !important;
    border-radius: 14px !important;
    box-shadow: 0 4px 15px rgba(6, 182, 212, 0.25) !important;
}
div[data-testid="stChatInput"] textarea:focus {
    border-color: #d946ef !important;
    box-shadow: 0 4px 20px rgba(217, 70, 239, 0.4) !important;
}
.stTextInput input {
    background-color: #1e293b !important;
    color: #f1f5f9 !important;
    border: 1px solid rgba(6, 182, 212, 0.3) !important;
    border-radius: 8px !important;
    box-shadow: 0 2px 8px rgba(6, 182, 212, 0.1) !important;
}
.stTextInput input:focus {
    border-color: #06b6d4 !important;
    box-shadow: 0 2px 12px rgba(6, 182, 212, 0.25) !important;
}

/* Fix input element labels to remain light in dark mode */
label[data-testid="stWidgetLabel"], .stSelectbox label, .stTextInput label, .stSlider label, .stCheckbox label {
    color: #f1f5f9 !important;
}
</style>
"""

# Helper to generate CSS for assistant chat message dynamically
def get_assistant_bubble_css(theme, color_name):
    if theme == "Rose Quartz Light":
        if color_name == "Championship Gold":
            bg = "linear-gradient(135deg, rgba(255, 250, 240, 0.95) 0%, rgba(253, 244, 219, 0.95) 100%)"
            border = "rgba(212, 175, 55, 0.3)"
            border_l = "#d4af37"
            shadow = "rgba(212, 175, 55, 0.1)"
        elif color_name == "Athletic Emerald":
            bg = "linear-gradient(135deg, rgba(240, 253, 244, 0.95) 0%, rgba(220, 252, 231, 0.95) 100%)"
            border = "rgba(16, 185, 129, 0.3)"
            border_l = "#10b981"
            shadow = "rgba(16, 185, 129, 0.1)"
        elif color_name == "Vibrant Pink":
            bg = "linear-gradient(135deg, rgba(255, 240, 245, 0.95) 0%, rgba(255, 228, 225, 0.95) 100%)"
            border = "rgba(255, 20, 147, 0.2)"
            border_l = "#ff1493"
            shadow = "rgba(255, 20, 147, 0.08)"
        else: # Neon Cyan
            bg = "linear-gradient(135deg, rgba(236, 254, 255, 0.95) 0%, rgba(207, 250, 254, 0.95) 100%)"
            border = "rgba(6, 182, 212, 0.3)"
            border_l = "#06b6d4"
            shadow = "rgba(6, 182, 212, 0.1)"
        text_color = "#000000"
    else: # Cyberpunk Dark
        if color_name == "Championship Gold":
            bg = "linear-gradient(135deg, rgba(15, 23, 42, 0.95) 0%, rgba(251, 191, 36, 0.04) 100%)"
            border = "rgba(251, 191, 36, 0.25)"
            border_l = "#fbbf24"
            shadow = "rgba(251, 191, 36, 0.12)"
        elif color_name == "Athletic Emerald":
            bg = "linear-gradient(135deg, rgba(15, 23, 42, 0.95) 0%, rgba(16, 185, 129, 0.04) 100%)"
            border = "rgba(16, 185, 129, 0.25)"
            border_l = "#10b981"
            shadow = "rgba(16, 185, 129, 0.12)"
        elif color_name == "Vibrant Pink":
            bg = "linear-gradient(135deg, rgba(15, 23, 42, 0.95) 0%, rgba(236, 72, 153, 0.04) 100%)"
            border = "rgba(236, 72, 153, 0.25)"
            border_l = "#ec4899"
            shadow = "rgba(236, 72, 153, 0.12)"
        else: # Neon Cyan
            bg = "linear-gradient(135deg, rgba(15, 23, 42, 0.95) 0%, rgba(6, 182, 212, 0.05) 100%)"
            border = "rgba(6, 182, 212, 0.2)"
            border_l = "#06b6d4"
            shadow = "rgba(6, 182, 212, 0.12)"
        text_color = "#f1f5f9"
        
    return f"""
    <style>
    .stChatMessage[data-testid="stChatMessageAssistant"] {{
        background: {bg} !important;
        border: 1px solid {border} !important;
        border-left: 5px solid {border_l} !important;
        box-shadow: 0 4px 20px {shadow} !important;
        color: {text_color} !important;
    }}
    </style>
    """

if st.session_state.current_theme == "Cyberpunk Dark":
    st.markdown(CYBERPUNK_DARK_CSS, unsafe_allow_html=True)
else:
    st.markdown(ROSE_QUARTZ_CSS, unsafe_allow_html=True)

# Dynamically inject assistant bubble styles
st.markdown(get_assistant_bubble_css(st.session_state.current_theme, st.session_state.ai_chat_color), unsafe_allow_html=True)

# Helper to automatically save API key to .env file
def save_key_to_env(provider_name, key_val):
    env_file_path = ".env"
    
    gemini_val = key_val if provider_name == "Google Gemini" else ENV_GEMINI_KEY
    anthropic_val = key_val if provider_name == "Anthropic Claude" else ENV_ANTHROPIC_KEY
    
    lines = []
    if os.path.exists(env_file_path):
        try:
            with open(env_file_path, "r", encoding="utf-8") as f:
                lines = f.readlines()
        except Exception:
            pass
            
    new_lines = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if "=" not in stripped:
            continue
        k = stripped.split("=", 1)[0].strip()
        if k in ["GEMINI_API_KEY", "GOOGLE_API_KEY", "ANTHROPIC_API_KEY"]:
            continue
        new_lines.append(line)
        
    if gemini_val:
        new_lines.append(f"GEMINI_API_KEY={gemini_val}\n")
    if anthropic_val:
        new_lines.append(f"ANTHROPIC_API_KEY={anthropic_val}\n")
        
    try:
        with open(env_file_path, "w", encoding="utf-8") as f:
            f.writelines(new_lines)
    except Exception as e:
        st.error(f"Unable to auto-save key to .env: {e}")

# ------------------------------------------------------------------
# Sidebar Settings
# ------------------------------------------------------------------
with st.sidebar:
    st.image("https://img.icons8.com/nolan/128/trophy.png", width=70)
    st.header("🏆 SportsAI Control Panel")

    # API Keys Configuration
    st.subheader("🔑 API Configuration")
    use_mock = st.toggle("Enable Demo / Mock Mode", value=False, help="Runs the app without needing API keys using offline simulation.")

    provider_options = []
    if HAS_GEMINI:
        provider_options.append("Google Gemini")
    if HAS_ANTHROPIC:
        provider_options.append("Anthropic Claude")
    if not provider_options:
        provider_options = ["Google Gemini", "Anthropic Claude"]

    provider = st.selectbox("AI Provider", provider_options, disabled=use_mock)

    # Load proper placeholder and value depending on choice
    if provider == "Google Gemini":
        env_key = ENV_GEMINI_KEY
        placeholder_text = "Enter Gemini API Key..."
    else:
        env_key = ENV_ANTHROPIC_KEY
        placeholder_text = "Enter Anthropic API Key..."

    api_key_input = st.text_input(
        f"{provider} API Key",
        value=env_key if not use_mock else "",
        type="password",
        placeholder=placeholder_text,
        disabled=use_mock,
        help="Provide your API key. If left empty, we will try to load it from the environmental .env variables."
    )

    # Automatically save key to .env when edited
    if api_key_input and api_key_input != env_key and not use_mock:
        save_key_to_env(provider, api_key_input)
        if provider == "Google Gemini":
            ENV_GEMINI_KEY = api_key_input
        else:
            ENV_ANTHROPIC_KEY = api_key_input
        st.toast(f"💾 Saved {provider} API Key to .env file!", icon="✅")

    active_api_key = api_key_input if api_key_input else env_key

    # Model Selection
    if provider == "Google Gemini":
        model_name = st.selectbox("Model", ["gemini-3.5-flash", "gemini-flash-latest", "gemini-2.0-flash", "gemini-2.5-pro", "gemini-3.1-flash-lite"], disabled=use_mock)
    else:
        model_name = st.selectbox("Model", ["claude-3-5-sonnet-latest", "claude-3-5-haiku-latest"], disabled=use_mock)

    st.markdown("---")

    st.subheader("🎨 Custom Theme")
    theme_choice = st.selectbox(
        "Theme Style",
        ["Rose Quartz Light", "Cyberpunk Dark"],
        index=0 if st.session_state.current_theme == "Rose Quartz Light" else 1
    )
    if theme_choice != st.session_state.current_theme:
        st.session_state.current_theme = theme_choice
        st.rerun()

    # Chat AI Theme color selection
    ai_color_choice = st.selectbox(
        "Chat AI Bubble Color",
        ["Championship Gold", "Athletic Emerald", "Vibrant Pink", "Neon Cyan"],
        index=["Championship Gold", "Athletic Emerald", "Vibrant Pink", "Neon Cyan"].index(st.session_state.ai_chat_color)
    )
    if ai_color_choice != st.session_state.ai_chat_color:
        st.session_state.ai_chat_color = ai_color_choice
        st.rerun()

    st.markdown("---")
    st.subheader("💬 Chat Sessions")
    
    # New chat button
    if st.button("➕ New Chat", use_container_width=True):
        new_id = f"chat_{int(time.time())}"
        st.session_state.chat_sessions[new_id] = {"name": f"New Chat {len(st.session_state.chat_sessions) + 1}", "messages": []}
        st.session_state.current_session_id = new_id
        st.session_state.messages = st.session_state.chat_sessions[new_id]["messages"]
        st.rerun()

    # List sessions
    for session_id in list(st.session_state.chat_sessions.keys()):
        session_info = st.session_state.chat_sessions[session_id]
        
        session_col1, session_col2 = st.columns([5, 1])
        with session_col1:
            label = session_info["name"]
            if session_id == st.session_state.current_session_id:
                label = f"👉 {label}"
            
            if st.button(label, key=f"select_{session_id}", use_container_width=True):
                st.session_state.current_session_id = session_id
                st.session_state.messages = st.session_state.chat_sessions[session_id]["messages"]
                st.rerun()
                
        with session_col2:
            is_delete_disabled = len(st.session_state.chat_sessions) <= 1
            if st.button("🗑️", key=f"delete_{session_id}", disabled=is_delete_disabled, help="Delete this session"):
                del st.session_state.chat_sessions[session_id]
                if session_id == st.session_state.current_session_id:
                    st.session_state.current_session_id = list(st.session_state.chat_sessions.keys())[0]
                st.session_state.messages = st.session_state.chat_sessions[st.session_state.current_session_id]["messages"]
                st.rerun()

    st.markdown("---")

    # User context preferences
    st.subheader("🎯 Athlete Profile")
    preferred_sport = st.selectbox("Primary Sport", ["Cricket", "Football", "Basketball", "Tennis", "Badminton", "Athletics", "Other"])
    skill_level = st.selectbox("Skill Level", ["Beginner", "Intermediate", "Advanced / Competitive"])
    target_goal = st.selectbox("Target Goal", ["General Knowledge", "Training Program", "Rules Clarification", "Fitness & Diet"])
    pref_language = st.selectbox("Language", ["English", "Hindi", "Kannada"])

    st.markdown("---")
    if st.button("🔄 Reset Conversational States"):
        st.session_state.chat_sessions = {"default": {"name": "Default Chat", "messages": []}}
        st.session_state.current_session_id = "default"
        st.session_state.messages = st.session_state.chat_sessions["default"]["messages"]
        st.session_state.rag_data = None
        st.session_state.quiz_started = False
        st.session_state.active_news_digest = None
        st.session_state.active_digest_source = None
        st.session_state.career_goals = []
        st.session_state.career_achievements = []
        st.session_state.career_chatbot_messages = []
        st.session_state.career_roadmap = None
        st.session_state.skill_assessment_results = None
        st.rerun()

    if active_api_key and not use_mock:
        st.success(f"🔑 {provider} Key Loaded Successfully!")
    elif not active_api_key and not use_mock:
        st.warning(f"⚠️ Configure {provider} key or enable 'Demo / Mock Mode' to start.")

# ------------------------------------------------------------------
# RAG Models Loader
# ------------------------------------------------------------------
@st.cache_resource
def get_embedding_model():
    if HAS_RAG:
        try:
            return SentenceTransformer("all-MiniLM-L6-v2")
        except Exception as e:
            st.error(f"Failed to load sentence-transformers model: {e}")
            return None
    return None

embedding_model = get_embedding_model()

# Helper to split text
def get_chunks(text, chunk_size=400, overlap=50):
    words = text.split()
    chunks = []
    for i in range(0, len(words), chunk_size - overlap):
        chunks.append(" ".join(words[i:i + chunk_size]))
    return chunks

# Export to docx
def generate_docx_plan(sport, level, goal, plan_content):
    if not HAS_DOCX:
        return None
    
    doc = docx.Document()
    doc.add_heading("SportsAI Personalized Schedule", level=0)
    
    p = doc.add_paragraph()
    p.add_run("Sport: ").bold = True
    p.add_run(f"{sport}\n")
    p.add_run("Skill level: ").bold = True
    p.add_run(f"{level}\n")
    p.add_run("Goal focus: ").bold = True
    p.add_run(f"{goal}\n")
    
    doc.add_heading("Your Custom Program Details", level=1)
    
    # Parse plan lines to preserve basic headings and bullet points
    for line in plan_content.split("\n"):
        line_stripped = line.strip()
        if not line_stripped:
            continue
        if line_stripped.startswith("###"):
            doc.add_heading(line_stripped.replace("###", "").strip(), level=3)
        elif line_stripped.startswith("##"):
            doc.add_heading(line_stripped.replace("##", "").strip(), level=2)
        elif line_stripped.startswith("#"):
            doc.add_heading(line_stripped.replace("#", "").strip(), level=1)
        elif line_stripped.startswith("-") or line_stripped.startswith("*"):
            doc.add_paragraph(line_stripped[1:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line_stripped)
            
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream

# ------------------------------------------------------------------
# Daily Sports News Helpers & Feed Config
# ------------------------------------------------------------------
import urllib.request
import xml.etree.ElementTree as ET
import re

NEWS_FEEDS = {
    "Top Headlines (BBC Sport)": "https://feeds.bbci.co.uk/sport/rss.xml",
    "Football News (BBC Sport)": "https://feeds.bbci.co.uk/sport/football/rss.xml",
    "ESPN Top Headlines": "https://www.espn.com/espn/rss/news",
    "ESPN NBA Headlines": "https://www.espn.com/espn/rss/nba/news",
    "ESPN NFL Headlines": "https://www.espn.com/espn/rss/nfl/news"
}

MOCK_NEWS = {
    "Top Headlines (BBC Sport)": [
        {
            "title": "🏆 Djokovic Clinches Historic 25th Grand Slam in Epic Final",
            "description": "In an absolute masterclass, Novak Djokovic wins his 25th Grand Slam title in straight sets, defying age and expectations in a thrilling championship match.",
            "link": "https://www.bbc.com/sport",
            "pubDate": "Sun, 09 Aug 2026 08:30:00 GMT"
        },
        {
            "title": "🏎️ Hamilton Takes Dramatic Victory at Wet Silverstone Grand Prix",
            "description": "Lewis Hamilton manages intermediate tires perfectly to claim a historic home victory at Silverstone after late rain triggers pitstop chaos.",
            "link": "https://www.bbc.com/sport",
            "pubDate": "Sun, 09 Aug 2026 07:15:00 GMT"
        },
        {
            "title": "🏏 India Secures Border-Gavaskar Trophy with Last-Over Heroics",
            "description": "A phenomenal partnership in the final hour seals the Test series for India, securing a historic win against Australia on a cracking day five pitch.",
            "link": "https://www.bbc.com/sport",
            "pubDate": "Sun, 09 Aug 2026 06:00:00 GMT"
        }
    ],
    "Football News (BBC Sport)": [
        {
            "title": "⚽ Real Madrid Crowned Champions of Europe Again After Dramatic Comeback",
            "description": "Two late goals from substitution forwards secure Real Madrid another UEFA Champions League title, defeating their opponents in a dramatic Wembley final.",
            "link": "https://www.bbc.com/sport/football",
            "pubDate": "Sun, 09 Aug 2026 08:00:00 GMT"
        },
        {
            "title": "🏆 Premier League Title Race Decided on Final Day Thriller",
            "description": "Manchester City holds off a fierce challenge from Arsenal, clinching the Premier League trophy with a commanding 3-1 victory on the final matchday.",
            "link": "https://www.bbc.com/sport/football",
            "pubDate": "Sun, 09 Aug 2026 07:45:00 GMT"
        },
        {
            "title": "🔄 Transfer Deadline Day: Record-Breaking Deals Done in Final Hours",
            "description": "A frantic close to the transfer window sees record-breaking sums exchanged as top clubs scramble to finalize squads before the midnight deadline.",
            "link": "https://www.bbc.com/sport/football",
            "pubDate": "Sat, 08 Aug 2026 23:30:00 GMT"
        }
    ],
    "ESPN Top Headlines": [
        {
            "title": "🥇 Opening Ceremony Lights Up the Summer Olympic Games",
            "description": "A spectacular flotilla along the river starts the Summer Olympic Games, showcasing culture, music, and athletic solidarity in front of millions of viewers.",
            "link": "https://www.espn.com",
            "pubDate": "Sun, 09 Aug 2026 05:00:00 GMT"
        },
        {
            "title": "🏌️ Masters Champion Dominates Augusta National to Claim Green Jacket",
            "description": "With a flawless final round, the world's top-ranked golfer secures a commanding victory at Augusta, finishing four strokes clear of the field.",
            "link": "https://www.espn.com",
            "pubDate": "Sat, 08 Aug 2026 20:30:00 GMT"
        }
    ],
    "ESPN NBA Headlines": [
        {
            "title": "🏀 Lakers Edge Celtics in Double-Overtime Classic at Madison Square Garden",
            "description": "A buzzer-beating three-pointer forces double overtime, where veteran stars lead a stunning run to secure a victory in the historic rivalry game.",
            "link": "https://www.espn.com/nba",
            "pubDate": "Sun, 09 Aug 2026 04:30:00 GMT"
        },
        {
            "title": "💎 Superstar Guards Sign Massive Max-Contract Extensions",
            "description": "Multiple All-Star players agree to record-setting extensions, locking in their futures and reshaping the salary-cap landscape of the league.",
            "link": "https://www.espn.com/nba",
            "pubDate": "Sat, 08 Aug 2026 19:15:00 GMT"
        }
    ],
    "ESPN NFL Headlines": [
        {
            "title": "🏈 Super Bowl MVP Engineers Historic Fourth-Quarter Comeback Drive",
            "description": "Trailing by ten with under five minutes remaining, the champion quarterback orchestrates a flawless 80-yard game-winning drive to lift the Vince Lombardi Trophy.",
            "link": "https://www.espn.com/nfl",
            "pubDate": "Sun, 09 Aug 2026 03:00:00 GMT"
        },
        {
            "title": "⏱️ Dynamic Draft Prospect Sets Combine Record with 4.21 Forty-Yard Dash",
            "description": "Scouts are left speechless as the rookie wide receiver breaks the all-time NFL Scouting Combine speed record, instantly boosting his draft value.",
            "link": "https://www.espn.com/nfl",
            "pubDate": "Sat, 08 Aug 2026 15:45:00 GMT"
        }
    ]
}

def fetch_sports_news(feed_name):
    # If in mock mode, immediately return local mock news
    if use_mock:
        return MOCK_NEWS.get(feed_name, [])
        
    feed_url = NEWS_FEEDS.get(feed_name)
    if not feed_url:
        return []
        
    try:
        req = urllib.request.Request(
            feed_url, 
            headers={'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/100.0.0.0 Safari/537.36'}
        )
        with urllib.request.urlopen(req, timeout=8) as response:
            xml_data = response.read()
            
        root = ET.fromstring(xml_data)
        items = []
        for item in root.findall('.//item')[:10]:  # Limit to 10 articles
            title_node = item.find('title')
            link_node = item.find('link')
            desc_node = item.find('description')
            pub_date_node = item.find('pubDate')
            
            title = title_node.text.strip() if title_node is not None and title_node.text else "No Title"
            link = link_node.text.strip() if link_node is not None and link_node.text else ""
            description = desc_node.text.strip() if desc_node is not None and desc_node.text else ""
            pub_date = pub_date_node.text.strip() if pub_date_node is not None and pub_date_node.text else ""
            
            # Clean up HTML tags if any in description (RSS descriptions sometimes contain HTML)
            description = re.sub('<[^<]+?>', '', description)
            
            items.append({
                "title": title,
                "link": link,
                "description": description,
                "pubDate": pub_date
            })
        return items
    except Exception as e:
        # Fallback to mock news on any network/parsing failure
        return MOCK_NEWS.get(feed_name, [])

# ------------------------------------------------------------------
# LLM Query Router
# ------------------------------------------------------------------
def call_llm(user_prompt, system_inst, history=None):
    if use_mock:
        # Mock mode responses
        if "training" in user_prompt.lower() or "plan" in user_prompt.lower() or "schedule" in user_prompt.lower():
            return f"""### 🏆 Weekly Training Schedule ({preferred_sport} - {skill_level})
This training plan focuses on **{target_goal}** to elevate your play style.

#### 📅 Weekly Breakdown
* **Monday: Aerobic Base & Core Focus**
  - 30 mins light jog/run.
  - 3 sets of plank (45s), Russian twists (20), and leg raises (15).
* **Tuesday: Technical Skills & Repetitive Drills**
  - {preferred_sport}-specific training. 40 minutes targeting basic footwork, stance, and hand-eye coordination.
* **Wednesday: Recovery & Mobility**
  - Active rest: 20 mins of deep full-body stretching and light walking.
* **Thursday: Strength & Endurance Power**
  - Squats (3 sets of 15), Push-ups (3 sets of 12), and lunges.
* **Friday: Match Simulation**
  - Play a light practice match or perform tactical scenarios under pressure.
* **Weekend: Rest & Rejuvenate**
  - Focus on nutritional intake and mental review.

*Disclaimer: This is a demo plan. Enable API keys to get tailored schedules.*
"""
        elif "news feed" in user_prompt.lower() or "headlines and summaries" in user_prompt.lower() or "sports news compiler" in user_prompt.lower():
            if "bbc" in user_prompt.lower() and "football" in user_prompt.lower():
                return """### ⚽ AI Daily Football Digest (Simulated)

* **Champions League Masterclass**: Real Madrid has done it again! Their dramatic comeback in the Wembley final demonstrates their unmatched resilience under pressure. Tactically, their mid-game adjustments in transition play proved decisive.
* **Premier League Finale**: Manchester City clinches their consecutive title, demonstrating incredible consistency over a grueling season. Arsenal's valiant chase sets up an exciting rivalry for the upcoming season.
* **Transfer Window Frenzy**: Club spending continues to soar. Teams are prioritizing highly versatile midfielders and dynamic box-to-box players to adapt to modern high-pressing systems.

> *"Football is a game of moments, and the teams that control transitions control the trophy."*"""
            elif "bbc" in user_prompt.lower():
                return """### 📰 AI Sports Headlines Digest (Simulated)

* **Tennis Milestone**: Novak Djokovic's 25th Grand Slam victory cements his legacy as the greatest of all time. His defensive adaptability and baseline coverage remain unmatched in long-format matches.
* **Silverstone Spectacle**: Hamilton's masterful drive in the wet conditions highlights the supreme importance of tire-wear management and strategic pitting communication in F1.
* **Border-Gavaskar Trophy**: India's tactical patience on a crumbling day-five pitch secures a historic series win. A masterclass in batting discipline and spinner rotation.

> *"Excellence is not a singular act, but a habit of continuous adaptation under pressure."*"""
            elif "nba" in user_prompt.lower():
                return """### 🏀 AI NBA News Digest (Simulated)

* **Overtime Thriller**: The Lakers-Celtics double-overtime matchup showcased elite isolation scoring and clutch defensive rotations. Small-ball lineups created significant space in the paint.
* **Financial Landscape**: Supermax contract extensions are locking in core rosters, forcing franchises to build depth through secondary trades and player development programs.

> *"The spacing and pace of the modern NBA require every position to be a threat from beyond the arc."*"""
            elif "nfl" in user_prompt.lower():
                return """### 🏈 AI NFL News Digest (Simulated)

* **Super Bowl Drive**: The championship-winning quarterback's fourth-quarter drive will be studied for years. Flawless execution of the two-minute drill against a hybrid zone-coverage scheme.
* **Draft Prospect Speed**: The record 4.21-second 40-yard dash has completely upended front-office draft boards, illustrating how elite vertical speed is valued in today's spread offenses.

> *"In the NFL, speed stretches defenses, but execution wins championships."*"""
            else:
                return """### 📰 AI Sports News General Digest (Simulated)

* **Olympic Excitement**: The spectacular opening ceremony has set a celebratory and unified tone for the global games. Keep an eye on track & field events where multiple world records are projected to fall.
* **Augusta Masterclass**: A flawless performance secures the green jacket, illustrating the value of precise iron-play and course management under high pressure.

> *"Sport has the power to change the world, to inspire, and to unite people like little else does."*"""
        elif "history" in user_prompt.lower() or "historic" in user_prompt.lower() or "milestone" in user_prompt.lower():
            sport = preferred_sport if 'preferred_sport' in locals() else "Sports"
            return f"""### ⏳ AI Sports History Digest ({sport}) - Simulated
Here is an overview of the historic growth and key evolutions in **{sport}**:

* **Origins & Early Days**: How the game transitioned from local pastime activities to organized, globally recognized governing bodies.
* **The Evolution of Rules**: Essential rule modifications (like the introduction of technology, VAR in football, DRS in cricket, or shot clocks in basketball) that revolutionized tactical depth and speed.
* **Legendary Icons**: Historical profiles of pioneering athletes who broke records and elevated the sport's global popularity.

To get a detailed deep-dive history of specific tournaments, matches, or legendary statistics, configure your API keys in the sidebar or run in live AI mode.

Would you like player statistics, match analysis, training plans, fitness advice, sports quizzes, tournament information, sports history, or career guidance? I'm ready to help!"""
        elif "career" in user_prompt.lower() or "guidance" in user_prompt.lower() or "roadmap" in user_prompt.lower() or "license" in user_prompt.lower():
            sport = preferred_sport if 'preferred_sport' in locals() else "Sports"
            return f"""### 🧭 AI Sports Career Roadmap - Simulated
Here is a simulated development pathway for a professional career in **{sport}**:

#### 📈 Key Phases of Development
1. **Foundation (Ages 6-12)**: Focus on building core motor skills, agility, and primary sport mechanics.
2. **Specialization (Ages 13-18)**: High-performance training, local league participation, and tactical sports intelligence.
3. **Professional Entry (Ages 18-22)**: Academy development, scouting trials, and sports agent linkages.
4. **Professional Transition**: Elite-level physical conditioning, sports nutrition plans, and mental coach guidance.

#### 🎓 Alternative Careers in Sports
* **Sports Analytics & Data Science**: Requires certification in python, data analysis, and tactical video tagging.
* **Coaching & Instruction**: Requires licensing paths (e.g., UEFA/AFC for Football, BCCI/ICC for Cricket, ITF for Tennis).
* **Sports Medicine**: Certified physiotherapist degrees or athletic trainer certificates.

To get a customized step-by-step career blueprint with specific licensing details and local academy paths, configure your API keys in the sidebar or run in live AI mode.

Would you like player statistics, match analysis, training plans, fitness advice, sports quizzes, tournament information, sports history, or career guidance? I'm ready to help!"""
        elif "quiz" in user_prompt.lower():
            return "Go to the interactive **🧩 Sports Quiz** tab to test your sports knowledge!"
        else:
            return f"👋 Hello! I am SportsAI (Demo Mode). I see you are interested in **{preferred_sport}** at the **{skill_level}** level. To unlock real AI responses from Anthropic Claude or Google Gemini, please add your API Key in the left Control Panel or configure the `.env` file!"

    if not active_api_key:
        return "⚠️ Please enter your API key or activate 'Demo / Mock Mode' in the sidebar control panel."

    full_system = f"{SYSTEM_PROMPT}\n\n[Active Athlete Profile]\nSport: {preferred_sport}; Level: {skill_level}; Goal: {target_goal}; Language: {pref_language}."

    # Gemini API Call
    if provider == "Google Gemini":
        if not HAS_GEMINI:
            return "⚠️ Gemini library is not installed."
        try:
            genai.configure(api_key=active_api_key)
            model = genai.GenerativeModel(model_name=model_name, system_instruction=full_system)
            
            gemini_messages = []
            if history:
                for h in history:
                    role = "user" if h["role"] == "user" else "model"
                    gemini_messages.append({"role": role, "parts": [h["content"]]})
            gemini_messages.append({"role": "user", "parts": [user_prompt]})
            
            response = model.generate_content(gemini_messages)
            return response.text
        except Exception as e:
            return f"⚠️ Google Gemini Error: {e}"

    # Anthropic Claude API Call
    else:
        if not HAS_ANTHROPIC:
            return "⚠️ Anthropic library is not installed."
        try:
            client = Anthropic(api_key=active_api_key)
            messages_list = []
            if history:
                for h in history:
                    messages_list.append({"role": h["role"], "content": h["content"]})
            messages_list.append({"role": "user", "content": user_prompt})

            response = client.messages.create(
                model=model_name,
                max_tokens=1500,
                system=full_system,
                messages=messages_list
            )
            return "".join([b.text for b in response.content if b.type == "text"])
        except Exception as e:
            return f"⚠️ Anthropic Claude Error: {e}"

# ------------------------------------------------------------------
# Sports History & Career Guidance Data Modules
# ------------------------------------------------------------------
HISTORICAL_TIMELINES = {
    "Cricket": [
        {"year": "1877", "title": "First Official Test Match", "desc": "Australia plays England at the Melbourne Cricket Ground, starting cricket's oldest rivalry."},
        {"year": "1909", "title": "Establishment of the ICC", "desc": "The Imperial Cricket Conference (now International Cricket Council) is founded to govern the sport globally."},
        {"year": "1975", "title": "First Men's Cricket World Cup", "desc": "West Indies defeats Australia in London to win the inaugural 60-over World Cup championship."},
        {"year": "2007", "title": "Inaugural T20 World Cup", "desc": "India defeats Pakistan by 5 runs in a nail-biting final, igniting the global T20 franchise revolution."},
        {"year": "2019", "title": "The Greatest ODI Final", "desc": "England defeats New Zealand on boundary countback after both the main match and super-over end in ties."}
    ],
    "Football (Soccer)": [
        {"year": "1863", "title": "Founding of the Football Association", "desc": "The Football Association (FA) meets in London, establishing the first unified, codified rules of soccer."},
        {"year": "1930", "title": "The First FIFA World Cup", "desc": "Uruguay hosts and wins the inaugural World Cup, defeating neighboring Argentina 4-2 in Montevideo."},
        {"year": "1970", "title": "Pele's Final World Cup Mastery", "desc": "Brazil wins their third World Cup in Mexico, introducing the 'Jogo Bonito' (beautiful game) to global TV screens."},
        {"year": "1992", "title": "Inception of the English Premier League", "desc": "The Premier League is established, revolutionizing football broadcasting rights, revenues, and global marketing."},
        {"year": "2022", "title": "Messi Cements Legacy in Qatar", "desc": "Lionel Messi leads Argentina to World Cup victory in one of the most dramatic finals in history against France."}
    ],
    "Basketball": [
        {"year": "1891", "title": "Dr. Naismith's Invention", "desc": "Dr. James Naismith invents basketball in Springfield, Massachusetts, using peach baskets and a soccer-style ball."},
        {"year": "1946", "title": "Establishment of the NBA", "desc": "The Basketball Association of America is founded, merging with the NBL in 1949 to form the National Basketball Association."},
        {"year": "1992", "title": "The Olympic 'Dream Team'", "desc": "NBA players compete in the Olympics for the first time, with Jordan, Johnson, and Bird capturing gold and global hearts."},
        {"year": "2016", "title": "Cavs Complete Historic Comeback", "desc": "Cleveland Cavaliers become the first team to overcome a 3-1 NBA Finals deficit to defeat the record-breaking Warriors."},
        {"year": "2023", "title": "LeBron James Breaks Scoring Record", "desc": "LeBron James surpasses Kareem Abdul-Jabbar's 38,387 points to become the NBA's all-time leading scorer."}
    ],
    "Tennis": [
        {"year": "1877", "title": "First Wimbledon Championship", "desc": "Wimbledon holds its inaugural tournament, starting the oldest and most prestigious tennis competition in history."},
        {"year": "1968", "title": "The Open Era Begins", "desc": "Grand Slam tournaments agree to allow professional players to compete alongside amateurs, elevating match quality and salaries."},
        {"year": "2008", "title": "Nadal-Federer Wimbledon Epic", "desc": "Rafael Nadal defeats Roger Federer in near-darkness in a 5-set final, widely hailed as the greatest tennis match ever played."},
        {"year": "2023", "title": "Djokovic Claims 24th Grand Slam", "desc": "Novak Djokovic wins the US Open to claim his 24th Major title, holding the outright record for men's tennis."},
        {"year": "2024", "title": "Golden Slam Milestones", "desc": "Novak Djokovic wins Olympic Gold in Paris, completing the Golden Slam and cementing his place among sporting legends."}
    ],
    "Badminton": [
        {"year": "1873", "title": "Badminton House Launch", "desc": "The Duke of Beaufort introduces the game 'badminton' to guests at his estate, adapting it from the Indian game Poona."},
        {"year": "1893", "title": "Badminton Association of England", "desc": "The world's first official governing body is formed, publishing the first official court dimensions and rules."},
        {"year": "1899", "title": "First All England Championships", "desc": "The inaugural All England Open is held, remaining the unofficial world championship for decades."},
        {"year": "1992", "title": "Olympic Medal Sport Debut", "desc": "Badminton is officially introduced as an Olympic medal sport at the Summer Games in Barcelona, Spain."},
        {"year": "2020", "title": "Axelsen Breaks Singaporean/Asian Streak", "desc": "Denmark's Viktor Axelsen wins Olympic Gold in Tokyo, showcasing dominant tactical play to break a long streak of Asian gold medalists."}
    ],
    "Olympics": [
        {"year": "1896", "title": "Modern Olympic Games Rebirth", "desc": "The first modern Olympic Games are held in Athens, Greece, with 241 athletes from 14 nations competing in 43 events."},
        {"year": "1936", "title": "Jesse Owens Defies Germany", "desc": "African-American sprinter Jesse Owens wins four gold medals in Berlin, breaking records and challenging discriminatory views."},
        {"year": "1972", "title": "Mark Spitz Wins Seven Golds", "desc": "Swimmer Mark Spitz sets seven world records and wins seven gold medals in a single Olympic Games in Munich."},
        {"year": "2008", "title": "Bolt & Phelps Dominance", "desc": "Usain Bolt shatters the 100m/200m world records, and Michael Phelps wins a historic 8 swimming gold medals in Beijing."},
        {"year": "2024", "title": "Urban Sports Rise in Paris", "desc": "Paris hosts the Olympic Games featuring street skate, breakdancing, and climbing, highlighting modern urban athletic culture."}
    ]
}

CAREERS_CATALOG = [
    {
        "title": "📊 Sports Data Analyst",
        "salary": "$65,000 - $120,000 / year",
        "desc": "Leverage statistics, coding (Python/R), and video tagging software to break down team tactics, evaluate player transfers, and optimize game-day performance.",
        "reqs": "Degree in Data Science, Statistics, or Math. Proficiency with SQL, Tableau, and specialized sports tracking tools (e.g. Opta, Wyscout)."
    },
    {
        "title": "⚽ Licensed Sports Coach",
        "salary": "$40,000 - $150,000+ / year",
        "desc": "Manage player development, design specialized training drills, organize team tactics, and direct squad coordination during competitive matches.",
        "reqs": "Coaching licenses issued by governing bodies (e.g., UEFA/AFC A/B/C licenses for soccer, BCCI level 1/2/3 for cricket, ITF certification)."
    },
    {
        "title": "💪 Strength & Conditioning Specialist",
        "salary": "$45,000 - $90,000 / year",
        "desc": "Design specialized physical conditioning, athletic speed, agility, and injury prevention programs for professional teams or academies.",
        "reqs": "BS in Exercise Science, Kinesiology, or Sports Science. Professional credentials like NSCA-CSCS or NASM-PES."
    },
    {
        "title": "🩹 Sports Physiotherapist",
        "salary": "$70,000 - $110,000 / year",
        "desc": "Diagnose, treat, and rehabilitate athletic injuries. Manage active recovery routines to optimize muscle health and return-to-play speed.",
        "reqs": "Doctor of Physical Therapy (DPT) degree. State board licensure, followed by Board Certification in Sports Physical Therapy (SCS)."
    },
    {
        "title": "💼 Sports Agent",
        "salary": "Commission-based (typically 3% - 10% of contracts)",
        "desc": "Negotiate professional playing contracts, structure lucrative sponsorship deals, and handle marketing, PR, and career management for athletes.",
        "reqs": "Degree in Sports Management, Business, or Law. Registration and certification with respective sport player unions (e.g., NBPA, NFLPA, FIFA)."
    },
    {
        "title": "🎙️ Sports Media & Journalism",
        "salary": "$35,000 - $85,000 / year",
        "desc": "Deliver live match play-by-play commentary, produce written analytical articles, or host television broadcasts for major sports networks.",
        "reqs": "Degree in Communications, Journalism, or Media studies. Strong portfolio of writing, audio reels, and on-air presentation skills."
    }
]

def render_timeline(sport_timeline, theme):
    primary_color = "#06b6d4" if theme == "Cyberpunk Dark" else "#ff1493"
    bg_color = "#0a0e17" if theme == "Cyberpunk Dark" else "#ffffff"
    text_color = "#cbd5e1" if theme == "Cyberpunk Dark" else "#333333"
    title_color = "#f1f5f9" if theme == "Cyberpunk Dark" else "#000000"
    
    html = f"""
    <style>
    .timeline-container {{
        padding: 0.5rem 1rem;
        position: relative;
        border-left: 2px dashed {primary_color};
        margin-left: 1.5rem;
    }}
    .timeline-event {{
        position: relative;
        margin-bottom: 1.8rem;
        padding-left: 1.5rem;
    }}
    .timeline-event::before {{
        content: '';
        position: absolute;
        left: -2.05rem;
        top: 0.25rem;
        width: 14px;
        height: 14px;
        border-radius: 50%;
        background-color: {primary_color};
        border: 2px solid {bg_color};
        box-shadow: 0 0 10px {primary_color}80;
    }}
    .timeline-year {{
        font-size: 1.15rem;
        font-weight: 700;
        color: {primary_color};
        margin: 0;
        line-height: 1.2;
    }}
    .timeline-title {{
        font-size: 1.05rem;
        font-weight: 600;
        color: {title_color};
        margin: 0.2rem 0;
    }}
    .timeline-desc {{
        font-size: 0.95rem;
        color: {text_color};
        margin: 0;
        line-height: 1.4;
    }}
    </style>
    <div class="timeline-container">
    """
    for item in sport_timeline:
        html += f"""
        <div class="timeline-event">
            <p class="timeline-year">{item['year']}</p>
            <h4 class="timeline-title">{item['title']}</h4>
            <p class="timeline-desc">{item['desc']}</p>
        </div>
        """
    html += "</div>"
    return html

# ------------------------------------------------------------------
# Main UI Presentation
# ------------------------------------------------------------------
st.markdown('<div class="main-title">🏆 SportsAI</div>', unsafe_allow_html=True)
st.markdown('<div class="subtitle">Your Dynamic Smart Coaching, Tactics, Rules & Training Platform</div>', unsafe_allow_html=True)

# Tabs
tab_chat, tab_planner, tab_news, tab_quiz, tab_history, tab_career, tab_rag, tab_about = st.tabs([
    "💬 Coaching Assistant", 
    "📋 Workout & Diet Planner", 
    "📰 Daily Sports News",
    "🧩 Sports Quiz Game", 
    "⏳ Sports History",
    "🧭 Career Guidance",
    "📂 Rulebook Semantic QA",
    "ℹ️ About SportsAI"
])

# ------------------------------------------------------------------
# Tab 1: Coaching Assistant Chat
# ------------------------------------------------------------------
with tab_chat:
    # Ensure messages is always bound to active session
    st.session_state.messages = st.session_state.chat_sessions[st.session_state.current_session_id]["messages"]

    # Display welcome card if history is empty
    if not st.session_state.messages:
        welcome_txt_color = "#cbd5e1" if st.session_state.current_theme == "Cyberpunk Dark" else "#333333"
        st.markdown(f"""
        <div class="welcome-card">
            <div class="welcome-header">Welcome to SportsAI Assistant! ⚡</div>
            <div style="color: {welcome_txt_color}; margin-bottom: 1.5rem; font-size: 1.05rem; line-height: 1.6;">
                I am your custom sports coach, rules expert, and tactical advisor. I can help you improve drills, analyze strategy, and answer rules questions. 
                Configure your preferred settings in the left sidebar.
            </div>
            <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(260px, 1fr)); gap: 1rem;">
                <div class="welcome-item"><span class="welcome-icon">⚡</span> Personal training & schedules</div>
                <div class="welcome-item"><span class="welcome-icon">⚡</span> Tactical analysis & comparisons</div>
                <div class="welcome-item"><span class="welcome-icon">⚡</span> Step-by-step game mechanics</div>
                <div class="welcome-item"><span class="welcome-icon">⚡</span> Multi-sport rules explanations</div>
            </div>
        </div>
        """, unsafe_allow_html=True)

        # Quick templates
        st.write("### 🚀 Quick Start Templates")
        cols = st.columns(3)
        prompts = [
            f"Provide 3 skill improvement drills for a {skill_level} {preferred_sport} player.",
            f"Explain the key rules of {preferred_sport} in simple terms.",
            f"Compare professional training regimes for {preferred_sport} athletes."
        ]
        for idx, prompt_text in enumerate(prompts):
            with cols[idx % 3]:
                if st.button(prompt_text, key=f"quick_prompt_{idx}"):
                    st.session_state.messages.append({"role": "user", "content": prompt_text})
                    with st.spinner("SportsAI is formulating a strategy..."):
                        reply = call_llm(prompt_text, SYSTEM_PROMPT, st.session_state.messages[:-1])
                        st.session_state.messages.append({"role": "assistant", "content": reply})
                    st.rerun()

    # Chat history
    for m in st.session_state.messages:
        with st.chat_message(m["role"]):
            st.markdown(m["content"])

    # Voice Input Section
    st.markdown('<div style="margin-top: 1.5rem; margin-bottom: 0.5rem; font-weight: 600; font-size: 0.95rem;">🎙️ Voice Input (Optional)</div>', unsafe_allow_html=True)
    voice_col1, voice_col2 = st.columns([1, 4])
    with voice_col1:
        voice_text = speech_to_text(
            language='en',
            start_prompt="🎙️ Start Speaking",
            stop_prompt="🛑 Stop & Process",
            just_once=True,
            key='voice_input'
        )
    with voice_col2:
        if voice_text:
            st.success(f"Recognized voice: **{voice_text}**")
        else:
            txt_help = "#cbd5e1" if st.session_state.current_theme == "Cyberpunk Dark" else "#666666"
            st.markdown(f'<span style="color: {txt_help}; font-size: 0.9rem; line-height: 2.2;">Click the mic button to speak your query.</span>', unsafe_allow_html=True)

    # Chat input
    chat_query = st.chat_input("Ask SportsAI about training, strategies, rules, or fitness...")
    
    # Process query if either typed or spoken
    active_query = None
    if chat_query:
        active_query = chat_query
    elif voice_text:
        active_query = voice_text
        
    if active_query:
        # Auto-rename chat session from default if first query
        curr_session = st.session_state.chat_sessions[st.session_state.current_session_id]
        if curr_session["name"] == "Default Chat" or curr_session["name"].startswith("New Chat"):
            clean_name = active_query.strip()
            if len(clean_name) > 22:
                clean_name = clean_name[:22] + "..."
            curr_session["name"] = f"💬 {clean_name}"

        st.session_state.messages.append({"role": "user", "content": active_query})
        with st.chat_message("user"):
            st.markdown(active_query)

        with st.chat_message("assistant"):
            with st.spinner("Analyzing strategy..."):
                reply = call_llm(active_query, SYSTEM_PROMPT, st.session_state.messages[:-1])
                st.markdown(reply)
        st.session_state.messages.append({"role": "assistant", "content": reply})
        st.rerun()

# ------------------------------------------------------------------
# Tab 2: Workout & Diet Planner (Exportable)
# ------------------------------------------------------------------
with tab_planner:
    st.subheader("📋 Athlete Weekly Schedule Generator")
    st.write("Generate a bespoke weekly training schedule and nutrition program customized to your profile details.")

    col_sport, col_level, col_goal = st.columns(3)
    with col_sport:
        planner_sport = st.selectbox("Select Sport", ["Cricket", "Football", "Basketball", "Tennis", "Badminton", "Running", "Swimming", "Weightlifting"], key="psport")
    with col_level:
        planner_level = st.selectbox("Your Level", ["Beginner", "Intermediate", "Advanced Elite Athlete"], key="plevel")
    with col_goal:
        planner_goal = st.selectbox("Focus Goal", ["Strength & Power", "Tactical & Skills Drill", "Cardiovascular Endurance", "Injury Rehab"], key="pgoal")

    col_time, col_diet = st.columns(2)
    with col_time:
        planner_time = st.select_slider("Daily Available Hours", options=["30 mins", "1 hour", "1.5 hours", "2+ hours"])
    with col_diet:
        planner_diet = st.checkbox("Include meal suggestions and macronutrient guidelines", value=True)

    if st.button("🛠️ Build My Training Plan", key="generate_plan_btn"):
        prompt = (
            f"Generate a detailed weekly training plan for a {planner_level} {planner_sport} player focusing on {planner_goal}. "
            f"The training should require about {planner_time} daily. "
        )
        if planner_diet:
            prompt += "Also include daily diet suggestions, hydration advice, and key macronutrient targets."

        with st.spinner("Compiling plan and formatting schedule..."):
            plan_output = call_llm(prompt, SYSTEM_PROMPT)
            st.session_state.active_plan = plan_output
            st.session_state.active_plan_sport = planner_sport
            st.session_state.active_plan_level = planner_level
            st.session_state.active_plan_goal = planner_goal

    # Display plan & export options if active plan exists
    if "active_plan" in st.session_state:
        st.markdown("---")
        st.markdown(st.session_state.active_plan)

        # Word document export
        if HAS_DOCX:
            word_stream = generate_docx_plan(
                st.session_state.active_plan_sport,
                st.session_state.active_plan_level,
                st.session_state.active_plan_goal,
                st.session_state.active_plan
            )
            if word_stream:
                st.download_button(
                    label="📥 Download Plan as Word Document (.docx)",
                    data=word_stream,
                    file_name=f"SportsAI_Training_{st.session_state.active_plan_sport}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.warning("⚠️ Word Document export disabled: 'python-docx' is not installed.")

# ------------------------------------------------------------------
# Tab 3: Daily Sports News
# ------------------------------------------------------------------
with tab_news:
    st.subheader("📰 Daily Sports News & AI Analysis")
    st.write("Browse real-time headlines from top sports publishers and generate automated AI digests, tactical summaries, and commentary.")
    
    col_source, col_fetch = st.columns([3, 1])
    with col_source:
        selected_source = st.selectbox("Select Sports News Feed", list(NEWS_FEEDS.keys()))
    with col_fetch:
        st.write("")
        st.write("")
        refresh_news = st.button("🔄 Refresh Feed", key="refresh_news_btn", use_container_width=True)
        
    with st.spinner("Fetching sports news..."):
        news_items = fetch_sports_news(selected_source)
        
    if not news_items:
        st.warning("No articles found in this feed or could not retrieve news. Try enabling Demo/Mock mode in the sidebar.")
    else:
        # AI news summary option
        st.write("")
        ai_col1, ai_col2 = st.columns([1, 1])
        with ai_col1:
            if st.button("🤖 Generate AI News Summary & Digest", key="ai_news_btn", use_container_width=True):
                # Prepare news items for LLM
                news_text = f"News Source: {selected_source}\n\n"
                for idx, item in enumerate(news_items[:6]):  # limit to top 6 for LLM token usage
                    news_text += f"{idx + 1}. {item['title']}\n"
                    news_text += f"   Date: {item['pubDate']}\n"
                    news_text += f"   Summary: {item['description']}\n\n"
                
                news_prompt = (
                    f"You are SportsAI, a premier sports commentator and analyst.\n"
                    f"Read the following recent news headlines and summaries, then construct a professional, exciting, "
                    f"and cohesive sports digest (3-4 concise bullet points of major events, tactical implications, "
                    f"and a quick, motivating quote or analysis style summary).\n\n"
                    f"Here is the news data:\n{news_text}"
                )
                
                with st.spinner("SportsAI is reviewing the match tape..."):
                    digest = call_llm(news_prompt, "You are a daily sports news compiler and commentator.")
                    st.session_state.active_news_digest = digest
                    st.session_state.active_digest_source = selected_source
                    st.rerun()
                    
        with ai_col2:
            if st.session_state.get("active_news_digest") and st.button("🗑️ Clear AI Digest", key="clear_digest_btn", use_container_width=True):
                st.session_state.active_news_digest = None
                st.session_state.active_digest_source = None
                st.rerun()

        # Display AI summary if it exists and matches current feed
        if st.session_state.get("active_news_digest") and st.session_state.get("active_digest_source") == selected_source:
            st.markdown("---")
            span_color = "#d946ef" if st.session_state.current_theme == "Cyberpunk Dark" else "#ff1493"
            txt_color = "#f1f5f9" if st.session_state.current_theme == "Cyberpunk Dark" else "#000000"
            st.markdown(f"""
            <div class="welcome-card" style="border-left: 5px solid {span_color};">
                <h3 style="color:{span_color}; margin-top:0;">🤖 AI News Digest & Tactical Analysis</h3>
                <div style="color:{txt_color}; font-size:1rem; line-height:1.6;">
            """, unsafe_allow_html=True)
            st.markdown(st.session_state.active_news_digest)
            st.markdown("""
                </div>
            </div>
            """, unsafe_allow_html=True)
            st.markdown("---")

        # News grid rendering
        st.write("### 📌 Latest Articles")
        
        # Display articles in cards
        for idx, item in enumerate(news_items):
            title = item.get("title", "No Title")
            description = item.get("description", "No description available.")
            link = item.get("link", "")
            pub_date = item.get("pubDate", "")
            
            # Formulate card content
            card_title_color = "#06b6d4" if st.session_state.current_theme == "Cyberpunk Dark" else "#ff1493"
            card_desc_color = "#cbd5e1" if st.session_state.current_theme == "Cyberpunk Dark" else "#444444"
            card_meta_color = "#94a3b8" if st.session_state.current_theme == "Cyberpunk Dark" else "#777777"
            
            st.markdown(f"""
            <div class="premium-card">
                <h4 style="margin: 0; color: {card_title_color}; font-size: 1.15rem; font-weight: 700;">{title}</h4>
                <p style="margin: 0.3rem 0; font-size: 0.85rem; color: {card_meta_color};">📅 Published: {pub_date}</p>
                <p style="margin: 0.5rem 0 1rem 0; font-size: 0.95rem; color: {card_desc_color}; line-height: 1.5;">{description}</p>
            </div>
            """, unsafe_allow_html=True)
            if link:
                st.markdown(f"[Read Full Article 🔗]({link})")
            st.markdown("<br>", unsafe_allow_html=True)

# ------------------------------------------------------------------
# Tab 4: Interactive Sports Quiz Game
# ------------------------------------------------------------------
with tab_quiz:
    st.subheader("🧩 Sports Trivia & GK Quiz")
    st.write("Test your sports rules, history, and records knowledge in an interactive game format.")

    # Quiz database
    LOCAL_QUIZZES = {
        "Cricket": [
            {
                "question": "How long is a standard cricket pitch between the wickets?",
                "options": ["20 Yards", "22 Yards", "24 Yards", "18 Yards"],
                "correct": 1,
                "explanation": "A standard cricket pitch stretches exactly 22 yards (20.12 meters) from stump to stump."
            },
            {
                "question": "Which bowler has taken the most wickets in Test Match history?",
                "options": ["Shane Warne", "Muttiah Muralitharan", "James Anderson", "Anil Kumble"],
                "correct": 1,
                "explanation": "Muttiah Muralitharan of Sri Lanka holds the record with 800 wickets in 133 Test matches."
            },
            {
                "question": "What is the maximum number of fielders allowed outside the 30-yard circle in powerplay 1 (first 10 overs) of a Men's ODI?",
                "options": ["2 fielders", "3 fielders", "4 fielders", "5 fielders"],
                "correct": 0,
                "explanation": "In Powerplay 1 of a standard ODI, only a maximum of 2 fielders are allowed outside the 30-yard circle."
            },
            {
                "question": "Which team won the inaugural ICC Men's T20 World Cup in 2007?",
                "options": ["Pakistan", "Australia", "India", "West Indies"],
                "correct": 2,
                "explanation": "India defeated Pakistan by 5 runs in a thrilling final in Johannesburg to win the first T20 World Cup."
            },
            {
                "question": "What does the abbreviation 'LBW' stand for in cricket rules?",
                "options": ["Leg Behind Wicket", "Leg Before Wicket", "Line Ball Wicket", "Low Bouncing Wicket"],
                "correct": 1,
                "explanation": "LBW stands for Leg Before Wicket, which is one of the methods of dismissing a batsman."
            }
        ],
        "Football (Soccer)": [
            {
                "question": "Which country has won the most FIFA Men's World Cups?",
                "options": ["Germany", "Italy", "Argentina", "Brazil"],
                "correct": 3,
                "explanation": "Brazil has won the World Cup 5 times (1958, 1962, 1970, 1994, 2002)."
            },
            {
                "question": "An direct free kick scored straight into the kicker's own goal results in what?",
                "options": ["A goal is awarded", "An indirect free kick", "A corner kick for the opponents", "A penalty kick"],
                "explanation": "If a free kick is kicked directly into a team's own goal, a corner kick is awarded to the opposing team to prevent own-goal exploits.",
                "correct": 2
            },
            {
                "question": "How long is the official duration of a standard professional association football match (excluding extra time)?",
                "options": ["80 minutes", "90 minutes", "100 minutes", "120 minutes"],
                "correct": 1,
                "explanation": "A standard match has two periods of 45 minutes each, totaling 90 minutes."
            },
            {
                "question": "Which player holds the record for most Ballon d'Or awards?",
                "options": ["Lionel Messi", "Cristiano Ronaldo", "Michel Platini", "Johan Cruyff"],
                "correct": 0,
                "explanation": "Lionel Messi has won the prestigious Ballon d'Or a record 8 times."
            },
            {
                "question": "What is the maximum number of players permitted on the field for one team during a match?",
                "options": ["10 players", "11 players", "12 players", "9 players"],
                "correct": 1,
                "explanation": "A match is played by two teams, each consisting of not more than 11 players; one must be the goalkeeper."
            }
        ],
        "General Sports": [
            {
                "question": "How many players are on the court for a single team in a basketball match?",
                "options": ["5 players", "6 players", "7 players", "4 players"],
                "correct": 0,
                "explanation": "A standard basketball game is played with 5 active players on court per team."
            },
            {
                "question": "Which sport uses a shuttlecock instead of a ball?",
                "options": ["Tennis", "Table Tennis", "Badminton", "Squash"],
                "correct": 2,
                "explanation": "Badminton is played using rackets to hit a shuttlecock across a high net."
            },
            {
                "question": "How often are the Olympic Games (Summer and Winter) held?",
                "options": ["Every 2 years", "Every 4 years", "Every 3 years", "Every 5 years"],
                "correct": 1,
                "explanation": "The Summer and Winter Olympics are each held every 4 years, staggered so there is an Olympic event every 2 years."
            },
            {
                "question": "What is the highest possible score in a single game of 10-pin bowling?",
                "options": ["200", "250", "300", "400"],
                "correct": 2,
                "explanation": "A perfect game of bowling consists of 12 consecutive strikes, resulting in a score of 300."
            },
            {
                "question": "Which grand slam tennis tournament is played on grass courts?",
                "options": ["Roland Garros", "Wimbledon", "US Open", "Australian Open"],
                "correct": 1,
                "explanation": "Wimbledon, held in London, is the only Grand Slam tennis tournament played on traditional grass courts."
            }
        ]
    }

    # Control Buttons
    col_sel, col_start = st.columns([3, 1])
    with col_sel:
        quiz_category = st.selectbox("Choose Quiz Subject", list(LOCAL_QUIZZES.keys()))
    with col_start:
        st.write("")
        st.write("")
        if st.button("🎬 Launch Quiz", key="launch_quiz_btn"):
            st.session_state.quiz_started = True
            st.session_state.quiz_questions = LOCAL_QUIZZES[quiz_category]
            st.session_state.quiz_current_index = 0
            st.session_state.quiz_score = 0
            st.session_state.quiz_answered = False
            st.session_state.quiz_selected_option = None
            st.rerun()

    # Quiz Execution
    if st.session_state.get("quiz_started", False):
        curr_idx = st.session_state.quiz_current_index
        questions = st.session_state.quiz_questions

        if curr_idx < len(questions):
            q_data = questions[curr_idx]
            
            # Display question box
            txt_color = "#f1f5f9" if st.session_state.current_theme == "Cyberpunk Dark" else "#000000"
            span_color = "#06b6d4" if st.session_state.current_theme == "Cyberpunk Dark" else "#ff1493"
            st.markdown(f"""
            <div class="premium-card">
                <span style="color:{span_color}; font-weight:600; font-size:0.9rem;">QUESTION {curr_idx + 1} OF {len(questions)}</span>
                <h3 style="margin-top:0.2rem; margin-bottom:1rem; color:{txt_color};">{q_data["question"]}</h3>
            </div>
            """, unsafe_allow_html=True)

            # Option selection using columns for a layout grid
            for opt_idx, option in enumerate(q_data["options"]):
                # Disable buttons after answering to prevent changing votes
                is_disabled = st.session_state.quiz_answered
                
                # Highlight selection color if answered
                btn_label = f"{chr(65+opt_idx)}. {option}"
                
                if st.button(btn_label, key=f"quiz_opt_{opt_idx}_{curr_idx}", disabled=is_disabled):
                    st.session_state.quiz_answered = True
                    st.session_state.quiz_selected_option = opt_idx
                    if opt_idx == q_data["correct"]:
                        st.session_state.quiz_score += 1
                    st.rerun()

            # Output results when answered
            if st.session_state.quiz_answered:
                selected = st.session_state.quiz_selected_option
                correct = q_data["correct"]
                
                st.markdown("---")
                if selected == correct:
                    st.success(f"🎉 **Correct!** Excellent work.")
                else:
                    st.error(f"❌ **Incorrect.** The correct answer was: **{q_data['options'][correct]}**")
                
                st.info(f"💡 **Explanation:** {q_data['explanation']}")
                
                if st.button("➡️ Proceed to Next Question", key="next_question_btn"):
                    st.session_state.quiz_current_index += 1
                    st.session_state.quiz_answered = False
                    st.session_state.quiz_selected_option = None
                    st.rerun()
        else:
            # Quiz complete!
            total = len(questions)
            score = st.session_state.quiz_score
            percentage = int((score / total) * 100)
            
            st.markdown("---")
            score_color = "#06b6d4" if st.session_state.current_theme == "Cyberpunk Dark" else "#ff1493"
            pct_color = "#cbd5e1" if st.session_state.current_theme == "Cyberpunk Dark" else "#333333"
            st.markdown(f"""
            <div class="welcome-card" style="text-align: center;">
                <h2>🏆 Quiz Completed!</h2>
                <h1 style="font-size:4rem; color:{score_color}; margin: 1rem 0;">{score} / {total}</h1>
                <p style="font-size:1.2rem; color:{pct_color};">You scored <b>{percentage}%</b></p>
            </div>
            """, unsafe_allow_html=True)
            
            if st.button("🔄 Play Again", key="play_again_btn"):
                st.session_state.quiz_started = False
                st.rerun()

# ------------------------------------------------------------------
# Tab 5: Sports History
# ------------------------------------------------------------------
with tab_history:
    st.subheader("⏳ Sports History & AI Explorer")
    st.write("Ask the AI to retrieve and narrate detailed history of matches, players, and rules.")

    st.write("### 🔍 AI History Explorer")
    st.write("Enter a historic match, legendary athlete, tournament, or rule evolution (e.g. *1983 Cricket World Cup Final*, *Pele*, *Evolution of Offside rule*):")
    
    hist_query = st.text_input("What would you like to explore?", placeholder="e.g. 1999 Champions League Final Manchester United vs Bayern Munich")
    
    if st.button("Explore History 🚀", key="explore_history_btn"):
        if hist_query:
            with st.spinner("Travelling back in time to retrieve sports archives..."):
                hist_prompt = (
                    f"You are SportsAI Pro acting as a Sports History Explorer.\n"
                    f"Write a highly engaging, detailed, and accurate historical summary of the following query: '{hist_query}'.\n"
                    f"Narrate in an exciting, storytelling tone. Include specific tactical details, key player contributions, "
                    f"relevant statistics, and the overall historical significance of this match, player, or rule evolution.\n\n"
                    f"End with the exact phrase:\n"
                    f"'Would you like player statistics, match analysis, training plans, fitness advice, sports quizzes, tournament information, sports history, or career guidance? I'm ready to help!'"
                )
                
                # Mock Mode responses for history
                if use_mock or not active_api_key:
                    query_lower = hist_query.lower()
                    if "1983" in query_lower:
                        hist_response = """### 🏏 The Miracle of Lord's: 1983 Cricket World Cup Final
India's victory in the 1983 Cricket World Cup final against the mighty West Indies is widely considered the most significant turning point in cricket history.

#### 📝 Match Summary
- **India**: 183 all out (54.4 overs) - Kris Srikkanth 38, Andy Roberts 3/32.
- **West Indies**: 140 all out (52 overs) - Viv Richards 33, Mohinder Amarnath 3/12, Madan Lal 3/31.
- **Result**: India won by 43 runs.

#### ⚔️ Tactical Breakdown
West Indies possessed the most fearsome four-pronged pace attack (Roberts, Garner, Marshall, Holding) and a batting lineup that had dominated the first two World Cups. After bowling India out for 183, West Indies expected a straightforward chase.
However, India's captain **Kapil Dev** inspired a legendary defensive fielding effort. The turning point was Kapil Dev's spectacular running catch to dismiss the devastating **Viv Richards** off Madan Lal's bowling. India's medium pacers bowled with immense discipline, exploiting the seam movement on a green Lord's wicket. Mohinder Amarnath's slow-medium cutters ripped through the lower order, securing him the Man of the Match.

#### 🌟 Historical Significance
This victory shattered the myth of West Indian invincibility and transformed cricket from an elite pastime to a national obsession in India, paving the way for the commercial and cultural powerhouse the sport is today.

Would you like player statistics, match analysis, training plans, fitness advice, sports quizzes, tournament information, sports history, or career guidance? I'm ready to help!"""
                    elif "offside" in query_lower:
                        hist_response = """### ⚽ The Tactical Evolution of the Offside Rule in Football
The offside rule has shaped the tactical landscape of association football more than any other regulation.

#### 📜 History of the Rule
- **1863**: The original FA rules stated that any player ahead of the ball was offside (similar to rugby).
- **1866**: The rule changed to require **three opponents** (usually the goalkeeper and two defenders) between the attacker and the goal when the ball was played.
- **1925**: To combat defensive stalemates, the rule was amended to require only **two opponents**. This triggered an immediate spike in goal-scoring and led to the creation of the famous "WM" formation by Arsenal manager Herbert Chapman.
- **1990**: The rule was adjusted to state that an attacker level with the second-last opponent is onside, encouraging attacking runs.
- **2005+ (Active/Passive)**: Players are only penalized if they are actively involved in play, interfering with an opponent, or gaining an advantage.

#### 📉 Tactical Impact
Modern tactics like Arrigo Sacchi's AC Milan "offside trap" of the late 1980s or Pep Guardiola's high defensive lines rely entirely on coordinated defensive movements to catch opponents offside.

Would you like player statistics, match analysis, training plans, fitness advice, sports quizzes, tournament information, sports history, or career guidance? I'm ready to help!"""
                    else:
                        hist_response = f"""### ⏳ AI Sports History Explorer: {hist_query} (Simulated)
This is an automated history brief for your query: **{hist_query}**.

* **The Event/Player**: Detailed records indicate this query represents a major landmark in sporting history, showcasing elite dedication, tactical excellence, and high-performance strategy.
* **Tactical Highlights**: Coaches and analysts of the era utilized creative set-play structures and advanced conditioning styles to counter their opponents.
* **Historical Legacy**: This moment continues to inspire modern coaching and athletic methodologies across international competitions.

*To unlock the full, highly detailed AI history narration powered by Anthropic Claude or Google Gemini, please configure your API keys in the sidebar control panel.*

Would you like player statistics, match analysis, training plans, fitness advice, sports quizzes, tournament information, sports history, or career guidance? I'm ready to help!"""
                else:
                    hist_response = call_llm(hist_prompt, "You are a Sports History Explorer.")
                
                st.session_state.history_explorer_result = hist_response
                st.rerun()

    if "history_explorer_result" in st.session_state:
        st.markdown("---")
        st.markdown(st.session_state.history_explorer_result)

# ------------------------------------------------------------------
# Tab 6: Career Guidance
# ------------------------------------------------------------------
with tab_career:
    st.subheader("🧭 Career Guidance & Athlete Pathways")
    st.write("Explore athletic development pathways, analyze your skills, find sports scholarships, track your progress, or consult our AI career counselor.")
    
    sub_career_tabs = st.tabs([
        "🎯 Career Roadmap & Builder",
        "🤖 Career AI Chatbot",
        "📊 Skill & Sport Selector",
        "🎓 Scholarships & Education",
        "💼 Alternative & Future Careers",
        "📈 Progress Tracker",
        "❓ Career FAQs"
    ])

    # Sub-tab 1: Career Roadmap
    with sub_career_tabs[0]:
        st.markdown("### 🎯 Personalized Career Roadmap Builder")
        st.write("Generate a customized developmental blueprint charting your path from grassroots training to elite sports competition and career entry.")
        
        c_col1, c_col2 = st.columns(2)
        with c_col1:
            r_age = st.slider("Your Current Age", min_value=6, max_value=50, value=16, key="r_age")
            r_sport = st.selectbox("Select Target Sport", ["Cricket", "Football", "Basketball", "Tennis", "Badminton", "Athletics", "Running", "Swimming"], key="r_sport")
        with c_col2:
            r_level = st.selectbox("Current Experience Level", ["Beginner (Grassroots)", "Intermediate (School / Club Competitor)", "Advanced (District / State / National level)"], key="r_level")
            r_goal = st.selectbox("Target Career Goal", ["Professional Athlete", "Licensed Sports Coach", "Sports Scientist / Tactical Analyst", "Sports Agent / Manager", "Referee / Match Official"], key="r_goal")
            
        if st.button("Generate Personalized Career Roadmap 🗺️", key="gen_roadmap_btn"):
            with st.spinner("Analyzing athlete attributes and mapping career pathways..."):
                roadmap_prompt = (
                    f"You are SportsAI Pro acting as a Sports Career Advisor.\n"
                    f"Generate a detailed, personalized sports career roadmap for a user who is {r_age} years old, plays {r_sport}, "
                    f"has a current level of '{r_level}', and aims to become a '{r_goal}'.\n"
                    f"Structure your response with clear headings:\n"
                    f"1. 📈 Career Recommendations & Core Fit (Evaluate suitability)\n"
                    f"2. 🧭 Career Pathway: Phase-by-Phase (Beginner → Training → Competition → Professional level)\n"
                    f"3. 🏫 Academy & Training Guidance (Suggest specific types of coaching, training volume, and academy setups)\n"
                    f"4. 🏆 Competition Information & Milestones (Explain levels of competition from school to national level needed to progress)\n"
                    f"5. 🎓 Sports Licensing & Certification Paths (Provide coaching or officiating license details if applicable, or professional requirements)\n\n"
                    f"Ensure the tone is motivating, realistic, and highly practical. End with the exact phrase:\n"
                    f"'Would you like player statistics, match analysis, training plans, fitness advice, sports quizzes, tournament information, sports history, or career guidance? I'm ready to help!'"
                )
                
                if use_mock or not active_api_key:
                    simulated_roadmap = f"""### 🗺️ Personalized Career Roadmap: {r_sport} ({r_goal})
**Prepared for**: {r_age}-year-old Athlete | **Level**: {r_level}

---

#### 1. 📈 Career Recommendations & Core Fit
Based on your age of {r_age} and interest in {r_sport}, pursuing a path towards becoming a **{r_goal}** is highly viable:
- **Fit Analysis**: At {r_age} years, starting or consolidating {r_sport} requires focused motor skill reinforcement or tactical specialization.
- **Alternative Fit**: If elite-level playing is not the sole goal, you can also transition into refereeing, coaching, or sports science.

---

#### 2. 🧭 Career Pathway: Grassroots to Professional
* **Phase 1: Foundation (Ages {r_age} to {max(r_age+2, 18)})**
  - Consolidate fundamental techniques. If beginner, dedicate 4-6 hours weekly to structural drills.
* **Phase 2: Training & Local Competition (Ages {max(r_age+2, 18)} to {max(r_age+5, 23)})**
  - Increase training volume to 8-12 hours per week. Move from school competitions to regional leagues.
* **Phase 3: High-Performance Selection (Ages {max(r_age+5, 23)}+)**
  - Participate in trials for state-level or franchise academies. Work with physical trainers and sports scientists to peak physically.

---

#### 3. 🏫 Academy & Training Guidance
- **Academy Recommendations**: Enroll in a certified {r_sport} academy (e.g. government sports authority or private club affiliated centers).
- **Training Frequency**: 4-5 sessions per week, focusing on technical skills (dribbling, shot precision, or footwork) and athletic conditioning.
- **Coaching Style**: Look for coaches with national certifications (e.g., BCCI Level 1 for cricket, AFC 'C' License for football, or equivalent sports science background).

---

#### 4. 🏆 Competition Information & Milestones
- **School & Club Level**: Dominate inter-school championships (e.g. SGFI) and district leagues.
- **State & National Level**: Focus on entering State Selection Trials to represent your state in national-level youth championships.
- **Professional Entry**: For professional sports, enter national league drafts, club trials, or collegiate division championships.

---

#### 5. 🎓 Sports Licensing & Certification Paths
- **For Coaching**: UEFA/AFC 'C', 'B', 'A' licensing courses or ICC/BCCI coaching levels.
- **For Officiating**: National federation referee/umpire exams (Category 3 to Category 1).
- **For Analytics/Management**: Certifications in sports data science (SQL/Python) or an MBA in Sports Management.

*This is a simulated roadmap. Configure your API keys in the control panel to generate a custom roadmap dynamically from our advanced AI model.*

Would you like player statistics, match analysis, training plans, fitness advice, sports quizzes, tournament information, sports history, or career guidance? I'm ready to help!"""
                    st.session_state.career_roadmap = simulated_roadmap
                else:
                    st.session_state.career_roadmap = call_llm(roadmap_prompt, "You are a Sports Career Advisor.")
                    
            st.rerun()
            
        if st.session_state.get("career_roadmap"):
            st.markdown("---")
            st.markdown(st.session_state.career_roadmap)
            
            st.download_button(
                label="📥 Download Roadmap (.txt)",
                data=st.session_state.career_roadmap,
                file_name=f"SportsAI_Career_Roadmap_{r_sport}.txt",
                mime="text/plain"
            )

    # Sub-tab 2: Career AI Chatbot
    with sub_career_tabs[1]:
        st.markdown("### 🤖 Sports Career AI Chatbot")
        st.write("Ask our specialized AI career counselor about scholarships, licensing courses (e.g. AFC, BCCI), academies, sports colleges, and other career paths.")
        
        for msg in st.session_state.career_chatbot_messages:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])
                
        if not st.session_state.career_chatbot_messages:
            with st.chat_message("assistant"):
                st.write("👋 Hello! I am your AI Sports Career Advisor. Ask me anything about sports academies, licensing certifications, sports scholarships, sports management, or how to balance sports with studies.")
        
        career_chat_input = st.chat_input("Ask about sports coaching, referee exams, scholarships, sports analytics...", key="career_chat_input_box")
        
        if career_chat_input:
            st.session_state.career_chatbot_messages.append({"role": "user", "content": career_chat_input})
            with st.chat_message("user"):
                st.markdown(career_chat_input)
                
            with st.chat_message("assistant"):
                with st.spinner("Consulting career guides..."):
                    c_chatbot_prompt = (
                        f"You are SportsAI Pro acting as a Sports Career Advisor.\n"
                        f"Answer the user's question about sports careers, academies, licensing (such as UEFA, AFC, BCCI, etc.), "
                        f"scholarships, sports science, or academics. Be professional, clear, and comprehensive.\n\n"
                        f"User Question: '{career_chat_input}'\n\n"
                        f"End with the exact phrase:\n"
                        f"'Would you like player statistics, match analysis, training plans, fitness advice, sports quizzes, tournament information, sports history, or career guidance? I'm ready to help!'"
                    )
                    
                    if use_mock or not active_api_key:
                        c_query = career_chat_input.lower()
                        if "scholarship" in c_query:
                            reply = """### 🎓 Sports Scholarships Guide
Here is some key guidance on sports scholarships:
- **India**: Under the **SAI (Sports Authority of India) Scheme** and **Khelo India Talent Development Scheme**, selected national-level junior athletes receive an annual scholarship of ₹6,20,000 (which includes ₹1,20,000 out-of-pocket allowance).
- **International (US NCAA)**: Universities in the US offer full/partial athletic scholarships for Division I and II programs. Eligibility is evaluated through tournament video tapes, national ranking portfolios, and SAT/ACT + TOEFL academic scores.
- **Corporate Schemes**: Reliance Foundation Youth Sports, TATA Football Academy, and JSW Sports run fully-funded residential programs covering training, gear, education, and travel.

Would you like player statistics, match analysis, training plans, fitness advice, sports quizzes, tournament information, sports history, or career guidance? I'm ready to help!"""
                        elif "license" in c_query or "coach" in c_query:
                            reply = """### ⚽ Coach Licensing Pathways
- **Football (Soccer)**: Governed by AFC/UEFA. The pathway is:
  1. **Grassroots / D-License**: Basic 3-6 day course on handling children.
  2. **C-License**: Focuses on youth team coaching (12-16 years).
  3. **B-License**: 3-week course on tactical systems and senior team coaching.
  4. **A-License**: Advanced coaching for professional club levels.
  5. **Pro-License**: Mandatory for coaching top-flight national teams and UEFA Champions League / AFC Champions League clubs.
- **Cricket**: Governed by ICC/BCCI:
  1. **BCCI Level 1**: Fundamentals of coaching junior and club cricketers.
  2. **BCCI Level 2**: Advanced tactical coaching, biomechanics of bowling/batting.
  3. **BCCI Level 3**: Elite coaching for Ranji Trophy, IPL, and national teams.

Would you like player statistics, match analysis, training plans, fitness advice, sports quizzes, tournament information, sports history, or career guidance? I'm ready to help!"""
                        else:
                            reply = f"""### 🧭 Career Guidance: {career_chat_input} (Simulated)
That is a great question! In professional sports, careers are highly competitive but very rewarding.
- **Suggested Pathway**: Look for certification programs with recognized federations, gather practical experience at local schools or clubs, and build a network by attending sports summits or webinars.
- **Next Steps**: Focus on building a portfolio. If you want to be a sports analyst, learn Python and SQL; if you want to coach, register for regional federation foundation courses.

*This is a simulated AI response. Configure your API keys in the sidebar to chat live with our advanced AI Sports Career Counselor.*

Would you like player statistics, match analysis, training plans, fitness advice, sports quizzes, tournament information, sports history, or career guidance? I'm ready to help!"""
                    else:
                        reply = call_llm(c_chatbot_prompt, "You are a Sports Career Advisor Chatbot.", st.session_state.career_chatbot_messages[:-1])
                        
                    st.markdown(reply)
            st.session_state.career_chatbot_messages.append({"role": "assistant", "content": reply})
            st.rerun()

    # Sub-tab 3: Skill & Sport Selector
    with sub_career_tabs[2]:
        st.markdown("### 📊 Skill Assessment & Sport Selector")
        st.write("Assess your physical and tactical attributes to find which sports align best with your skills and interests.")
        
        col_s1, col_s2 = st.columns(2)
        with col_s1:
            st.write("##### 🏃 Physical & Tactical Attributes")
            sk_speed = st.slider("Speed & Agility (Sprinting, acceleration, change of direction)", 1, 10, 5, key="sk_speed")
            sk_stamina = st.slider("Stamina & Endurance (Aerobic capacity, running long distances)", 1, 10, 5, key="sk_stamina")
            sk_coord = st.slider("Coordination & Reflexes (Hand-eye, foot-eye coordination, reaction time)", 1, 10, 5, key="sk_coord")
            sk_strength = st.slider("Strength & Power (Explosive jumps, throws, lifting capacity)", 1, 10, 5, key="sk_strength")
        with col_s2:
            st.write("##### 🤝 Teamwork & Mental Attributes")
            sk_team = st.slider("Teamwork & Communication (Playing in groups, leadership, strategy sharing)", 1, 10, 5, key="sk_team")
            sk_tech = st.slider("Technique & Tactical IQ (Understanding rules, spacing, and play execution)", 1, 10, 5, key="sk_tech")
            
            st.write("##### 🎯 Environment & Style Preferences")
            sk_env = st.selectbox("Preferred Playing Environment", ["Outdoor", "Indoor", "Water (Swimming)", "No Preference"], key="sk_env")
            sk_contact = st.selectbox("Physical Contact Level", ["No Contact (e.g. Tennis, Athletics)", "Limited Contact (e.g. Cricket, Basketball)", "High Contact (e.g. Football, Rugby, Kabaddi)"], key="sk_contact")
            sk_format = st.selectbox("Game Format", ["Team Sport", "Individual Sport", "No Preference"], key="sk_format")
            
        if st.button("Analyze Skills & Match Sports 🔍", key="analyze_skills_btn"):
            scores = {}
            
            # Cricket
            cricket_score = (sk_coord * 3 + sk_team * 2 + sk_tech * 2 + sk_speed * 1) / 8
            if sk_contact == "High Contact": cricket_score -= 2
            if sk_env == "Water (Swimming)": cricket_score -= 4
            if sk_format == "Individual Sport": cricket_score -= 3
            scores["Cricket"] = max(0.0, cricket_score)
            
            # Football
            football_score = (sk_stamina * 3 + sk_speed * 2 + sk_team * 2 + sk_coord * 1 + sk_tech * 2) / 10
            if sk_contact == "No Contact": football_score -= 2
            if sk_env == "Water (Swimming)" or sk_env == "Indoor": football_score -= 3
            if sk_format == "Individual Sport": football_score -= 3
            scores["Football (Soccer)"] = max(0.0, football_score)
            
            # Athletics
            athletics_score = (sk_speed * 4 + sk_stamina * 3 + sk_strength * 3) / 10
            if sk_format == "Team Sport": athletics_score -= 3
            if sk_contact == "High Contact": athletics_score -= 1
            scores["Athletics (Sprinting/Long Distance)"] = max(0.0, athletics_score)
            
            # Badminton
            badminton_score = (sk_coord * 4 + sk_speed * 3 + sk_stamina * 2 + sk_tech * 2) / 11
            if sk_contact == "High Contact": badminton_score -= 3
            if sk_env == "Outdoor": badminton_score -= 1
            if sk_format == "Team Sport": badminton_score -= 2
            scores["Badminton"] = max(0.0, badminton_score)
            
            # Basketball
            basketball_score = (sk_speed * 3 + sk_coord * 3 + sk_team * 2 + sk_stamina * 2 + sk_strength * 1) / 11
            if sk_contact == "No Contact": basketball_score -= 3
            if sk_env == "Water (Swimming)": basketball_score -= 4
            if sk_format == "Individual Sport": basketball_score -= 3
            scores["Basketball"] = max(0.0, basketball_score)
            
            sorted_sports = sorted(scores.items(), key=lambda x: x[1], reverse=True)
            
            st.session_state.skill_assessment_results = {
                "sports": sorted_sports,
                "skills": {
                    "Speed": sk_speed, "Stamina": sk_stamina, "Coordination": sk_coord,
                    "Strength": sk_strength, "Teamwork": sk_team, "Technique": sk_tech
                },
                "preferences": {"env": sk_env, "contact": sk_contact, "format": sk_format}
            }
            st.rerun()
            
        if st.session_state.get("skill_assessment_results"):
            results = st.session_state.skill_assessment_results
            st.markdown("---")
            st.subheader("🎯 Assessment Results & Recommendations")
            
            st.write("##### 🏆 Best Suited Sports for You:")
            cols_r = st.columns(len(results["sports"]))
            for idx, (sport, score) in enumerate(results["sports"]):
                with cols_r[idx]:
                    percentage = int((score / 10) * 100)
                    st.metric(label=sport, value=f"{percentage}% Match")
                    
            primary_sport = results["sports"][0][0]
            st.markdown(f"""
            <div class="welcome-card" style="border-left: 5px solid #10b981;">
                <h4 style="color:#10b981; margin-top:0;">Sport Selection Guidance: {primary_sport}</h4>
                <p>Based on your physical strengths (especially <b>{", ".join([k for k, v in results["skills"].items() if v >= 7]) or "balanced attributes"}</b>), 
                you show high potential in <b>{primary_sport}</b>. Your preferred contact level ({results["preferences"]["contact"]}) and environment ({results["preferences"]["env"]}) also align well.</p>
                <b>Recommended Career Pathways in this Sport:</b>
                <ul>
                    <li><b>Professional Athlete</b>: Compete in state, national, or franchise leagues.</li>
                    <li><b>High Performance Coach</b>: Leverage your tactical understanding to train upcoming teams.</li>
                    <li><b>Sports Analyst</b>: Use technical statistics and match videos to analyze team plays.</li>
                </ul>
            </div>
            """, unsafe_allow_html=True)
            
            st.write("##### 🏋️ Age-Appropriate Fitness Guidance")
            athlete_age = st.session_state.get("r_age", 16)
            
            if athlete_age < 12:
                fitness_tip = """
                - **Focus**: Fundamental motor skills, balance, coordination, and agility.
                - **Activity**: Active play, gymnastics, running games, light swimming. Avoid intensive weight training.
                - **Volume**: 60 minutes of active play daily. Keep it fun and multi-sport oriented to avoid repetitive stress injuries.
                """
            elif 12 <= athlete_age <= 18:
                fitness_tip = """
                - **Focus**: Agility, speed drills, flexibility, and foundational strength.
                - **Activity**: Bodyweight exercises (pushups, pullups, planks), plyometrics, light resistance training with proper form, and specific interval running.
                - **Volume**: 4-5 days/week. Prioritize learning correct lifting techniques under certified supervision before adding heavy loads.
                """
            elif 19 <= athlete_age <= 35:
                fitness_tip = """
                - **Focus**: Power development, periodized strength training, anaerobic and aerobic capacity.
                - **Activity**: Olympic lifts, compound lifts (squats, deadlifts, bench press), high-intensity interval training (HIIT), and sport-specific tactical training.
                - **Volume**: 5-6 days/week. Use periodization (macrocycles/microcycles) to peak during competitive match seasons.
                """
            else:
                fitness_tip = """
                - **Focus**: Active recovery, joint mobility, core stability, and cardiovascular health.
                - **Activity**: Low-impact aerobics, yoga, swimming, moderate resistance training.
                - **Volume**: 3-4 days/week. Focus on functional movement patterns to maintain muscle density and prevent joints from stiffening.
                """
            st.info(f"💡 **Fitness Advice (Suggested for Age {athlete_age}):** {fitness_tip}")

    # Sub-tab 4: Scholarships & Education
    with sub_career_tabs[3]:
        st.markdown("### 🎓 Sports Scholarships & Academic Courses")
        st.write("Explore financial aid options and academic pathways that help you balance sports excellence with professional qualifications.")
        
        tab_edu1, tab_edu2 = st.tabs(["🏆 Sports Scholarships", "📖 Sports Education Courses"])
        
        with tab_edu1:
            st.write("##### 💰 Major Sports Scholarship Programs")
            
            scholarships = [
                {
                    "name": "🇮🇳 SAI National Sports Promotional Schemes",
                    "provider": "Sports Authority of India (Government)",
                    "eligibility": "Sub-junior, junior, and senior athletes who have won medals at State/National championships.",
                    "details": "Covers fully funded boarding, food, equipment, medical facilities, coaching, and kit allowance at SAI training centers."
                },
                {
                    "name": "🔥 Khelo India Talent Development Scholarship",
                    "provider": "Ministry of Youth Affairs & Sports, India",
                    "eligibility": "Junior athletes identified through the Khelo India Games or national scouting committees.",
                    "details": "Financial assistance of ₹6,20,000 per annum (₹1,20,000 out-of-pocket allowance, ₹5,00,000 for training, nutrition, and academy fees)."
                },
                {
                    "name": "🇺🇸 US College Sports Scholarships (NCAA Division I & II)",
                    "provider": "NCAA Universities (United States)",
                    "eligibility": "Outstanding high school athletes with high national rankings and SAT/TOEFL scores.",
                    "details": "Full or partial tuition fee waiver, hostel boarding, training equipment, and international competition representation."
                },
                {
                    "name": "🌐 Reliance Foundation Youth Sports Scholarships",
                    "provider": "Reliance Foundation (Corporate)",
                    "eligibility": "Top-tier student-athletes selected from RFYS school and college athletics, football, and cricket championships.",
                    "details": "Covers tuition fees support, high-performance training camps, nutritional consultations, and professional academy access."
                }
            ]
            
            for s in scholarships:
                st.markdown(f"""
                <div class="premium-card">
                    <h4 style="margin: 0; color: #ff1493;">{s["name"]}</h4>
                    <p style="margin: 0.3rem 0; font-size: 0.9rem; font-weight: bold; color: #06b6d4;">Provider: {s["provider"]}</p>
                    <p style="margin: 0.5rem 0; font-size: 0.95rem;"><b>Eligibility Criteria:</b> {s["eligibility"]}</p>
                    <p style="margin: 0.5rem 0; font-size: 0.95rem; line-height: 1.5;"><b>Award & Benefits:</b> {s["details"]}</p>
                </div>
                """, unsafe_allow_html=True)
                
        with tab_edu2:
            st.write("##### 🎓 Sports-Related Professional Courses")
            st.write("If you want to build a career in the sports industry alongside or after your playing days, explore these specialized programs:")
            
            courses = [
                {
                    "degree": "1. BS / MS in Sports Science & Kinesiology",
                    "duration": "3 - 4 Years (UG) / 2 Years (PG)",
                    "scope": "Study human movement, sports biomechanics, exercise physiology, and athletic training mechanics.",
                    "careers": "High-Performance Analyst, Strength & Conditioning Coach, Biomechanist."
                },
                {
                    "degree": "2. MBA / PGDM in Sports Management",
                    "duration": "2 Years",
                    "scope": "Study sports marketing, sponsorship negotiation, stadium management, event logistics, and player agency operations.",
                    "careers": "Sports Agent, Tournament Director, Marketing Head, Stadium Manager."
                },
                {
                    "degree": "3. Bachelor of Physical Education (B.P.Ed / M.P.Ed)",
                    "duration": "2 - 4 Years",
                    "scope": "Study physical teaching methodologies, sports psychology, safety, and school coaching systems.",
                    "careers": "School Sports Director, Physical Education Teacher, Youth Coordinator."
                },
                {
                    "degree": "4. Diploma in Sports Coaching (NSNIS)",
                    "duration": "1 Year",
                    "scope": "Governed by Netaji Subhas National Institute of Sports (India). Intensive coaching training in a specialized sport.",
                    "careers": "Government-appointed Coach, Elite Academy Trainer."
                }
            ]
            
            for c in courses:
                st.markdown(f"""
                <div class="premium-card">
                    <h4 style="margin: 0; color: #06b6d4;">{c["degree"]}</h4>
                    <p style="margin: 0.3rem 0; font-size: 0.9rem; font-weight: bold; color: #ff1493;">Duration: {c["duration"]}</p>
                    <p style="margin: 0.5rem 0; font-size: 0.95rem;"><b>Course Focus:</b> {c["scope"]}</p>
                    <p style="margin: 0.5rem 0; font-size: 0.95rem;"><b>Key Careers:</b> {c["careers"]}</p>
                </div>
                """, unsafe_allow_html=True)
                
            st.markdown("""
            <div class="welcome-card" style="margin-top:1.5rem;">
                <h4 style="color:#ff1493; margin-top:0;">📚 Balancing Academics & Sports</h4>
                <p style="line-height:1.6; font-size:0.95rem;">
                    1. 📅 <b>Time Blocking</b>: Segment your day strictly. Dedicate 5:00 AM - 8:00 AM for training, 9:00 AM - 3:00 PM for school/college classes, and 5:00 PM - 7:00 PM for tactical reviews/strength work.
                    <br>2. 🏫 <b>Sports Quota Admissions</b>: Most universities have reserved seats and relaxed grade cut-offs for state and national medal-winning athletes. Check guidelines before applying.
                    <br>3. 🎒 <b>Flexible Learning</b>: Consider enrolling in distance-learning or online high schools if international travel for tournaments requires frequent absences.
                </p>
            </div>
            """, unsafe_allow_html=True)

    # Sub-tab 5: Alternative & Future Careers
    with sub_career_tabs[4]:
        st.markdown("### 💼 Alternative & Future Sports Careers")
        st.write("Understand licensing pathways for traditional sports careers and explore the rapid growth of technology-driven future careers.")
        
        tab_c1, tab_c2 = st.tabs(["💼 Traditional Alternative Careers", "🚀 Emerging Future Fields"])
        
        with tab_c1:
            st.write("##### 🛠️ Licensing and Entry Paths for Key Roles")
            
            alt_careers = [
                {
                    "title": "⚽ Sports Coach / Instructor",
                    "salary": "$40,000 - $120,000+",
                    "path": "Must obtain licensing from respective national or international federations (e.g. AFC D/C/B/A coaching license for soccer, BCCI Level 1/2/3 for cricket, ITF for tennis)."
                },
                {
                    "title": "🏁 Referee / Umpire / Match Official",
                    "salary": "$30,000 - $100,000+ (per match fee basis)",
                    "path": "Attend local/state referee clinics, pass physical fitness tests (beep tests) and theoretical rules exams. Advance from Category-3 (District) to Category-1 (National) and finally FIFA / ICC International panels."
                },
                {
                    "title": "🩹 Sports Physiotherapist",
                    "salary": "$60,000 - $110,000",
                    "path": "Complete a Bachelor of Physiotherapy (BPT) followed by a Master of Physiotherapy (MPT) in Sports. Obtain registrations with state medical boards."
                },
                {
                    "title": "🎙️ Sports Journalist / Broadcaster",
                    "salary": "$35,000 - $80,000",
                    "path": "Earn a degree in Journalism, Mass Communication, or English. Build a portfolio of written match reviews, podcast episodes, or video analysis scripts."
                }
            ]
            
            for ac in alt_careers:
                st.markdown(f"""
                <div class="premium-card">
                    <h4 style="margin: 0; color: #06b6d4;">{ac["title"]}</h4>
                    <p style="margin: 0.3rem 0; font-size: 0.9rem; font-weight: bold; color: #ff1493;">Average Salary: {ac["salary"]}</p>
                    <p style="margin: 0.5rem 0; font-size: 0.95rem; line-height: 1.5;"><b>Licensing & Entry Pathway:</b> {ac["path"]}</p>
                </div>
                """, unsafe_allow_html=True)
                
        with tab_c2:
            st.write("##### 🤖 Emerging Fields in Sports Technology & Analytics")
            st.write("Explore careers at the intersection of sports, engineering, data science, and AI:")
            
            future_fields = [
                {
                    "title": "📊 Sports Analytics & Data Science",
                    "skills": "Python, SQL, Tableau, R, machine learning algorithms, Wyscout/Opta database queries.",
                    "role": "Evaluate player performance using metrics like expected goals (xG), match data coding, and transfer window statistics to optimize squad build decisions."
                },
                {
                    "title": "🔌 Sports Technology & Wearable Engineering",
                    "skills": "Hardware engineering, biosensors, signal processing, IoT integration.",
                    "role": "Design smart vests (e.g. GPS trackers like Catapult), heart rate monitors, and smart equipment (Hawkeye sensor-embedded bats/balls) to track real-time physical workload."
                },
                {
                    "title": "🤖 Artificial Intelligence in Sports",
                    "skills": "Computer vision, deep learning, PyTorch/TensorFlow, video streaming pipelines.",
                    "role": "Develop automated camera tracking systems (e.g. automated soccer coaching highlights, ball-tracking algorithms, automated line-calls, and injury prediction algorithms based on muscle load)."
                },
                {
                    "title": "💼 Sports Franchise & Brand Management",
                    "skills": "Sponsorship valuation, social media analytics, player image rights, contract negotiation.",
                    "role": "Manage brand representation for elite leagues (like IPL, Premier League, NBA) or esports organizations, maximizing commercial engagement."
                }
            ]
            
            for ff in future_fields:
                st.markdown(f"""
                <div class="premium-card">
                    <h4 style="margin: 0; color: #ff1493;">{ff["title"]}</h4>
                    <p style="margin: 0.3rem 0; font-size: 0.9rem; font-weight: bold; color: #06b6d4;">Key Technical Skills: {ff["skills"]}</p>
                    <p style="margin: 0.5rem 0; font-size: 0.95rem; line-height: 1.5;"><b>Role Description:</b> {ff["role"]}</p>
                </div>
                """, unsafe_allow_html=True)

    # Sub-tab 6: Progress Tracker
    with sub_career_tabs[5]:
        st.markdown("### 📈 Goal Planning & Competition Progress Tracker")
        st.write("Track your training milestones, log competitive participations, and monitor your athletic progress over time.")
        
        completed_goals = sum(1 for g in st.session_state.career_goals if g["status"] == "Completed")
        total_goals = len(st.session_state.career_goals)
        total_comps = len(st.session_state.career_achievements)
        
        st.write("##### 📊 Your Athlete Dashboard")
        d_col1, d_col2, d_col3 = st.columns(3)
        with d_col1:
            st.metric(label="Competitions Logged", value=total_comps)
        with d_col2:
            st.metric(label="Goals Completed", value=f"{completed_goals} / {total_goals}")
        with d_col3:
            progress_pct = int((completed_goals / total_goals) * 100) if total_goals > 0 else 0
            st.metric(label="Goal Completion Rate", value=f"{progress_pct}%")
            
        if total_goals > 0:
            st.progress(completed_goals / total_goals)
            
        st.markdown("---")
        
        t_col1, t_col2 = st.columns(2)
        with t_col1:
            st.write("##### 🎯 Create a Training Goal")
            g_title = st.text_input("Goal Title (e.g. Practice 50 penalties, run 5km under 20 mins)", key="g_title_input")
            g_date = st.date_input("Target Completion Date", key="g_date_input")
            g_status = st.selectbox("Status", ["Not Started", "In Progress", "Completed"], key="g_status_input")
            
            if st.button("➕ Add Goal", key="add_goal_btn", use_container_width=True):
                if g_title:
                    st.session_state.career_goals.append({
                        "title": g_title,
                        "date": str(g_date),
                        "status": g_status
                    })
                    st.success(f"Added goal: '{g_title}'!")
                    st.rerun()
                else:
                    st.warning("Please enter a goal title.")
                    
        with t_col2:
            st.write("##### 🏆 Log Competition Participation")
            comp_name = st.text_input("Tournament Name", key="comp_name_input")
            comp_sport = st.selectbox("Sport", ["Cricket", "Football", "Basketball", "Tennis", "Badminton", "Athletics", "Other"], key="comp_sport_input")
            comp_level = st.selectbox("Competition Level", ["School / College", "District / Club", "State / Zonal", "National", "International"], key="comp_level_input")
            comp_date = st.date_input("Competition Date", key="comp_date_input")
            comp_result = st.text_input("Result / Achievement (e.g. Winner, Gold Medal, Quarter Finals)", key="comp_result_input")
            
            if st.button("➕ Log Competition", key="log_comp_btn", use_container_width=True):
                if comp_name and comp_result:
                    st.session_state.career_achievements.append({
                        "event": comp_name,
                        "sport": comp_sport,
                        "level": comp_level,
                        "date": str(comp_date),
                        "result": comp_result
                    })
                    st.success(f"Logged achievement for '{comp_name}'!")
                    st.rerun()
                else:
                    st.warning("Please fill in the Tournament Name and Result.")
                    
        st.markdown("---")
        
        l_col1, l_col2 = st.columns(2)
        with l_col1:
            st.write("##### 🎯 Active Goals List")
            if not st.session_state.career_goals:
                st.write("No goals set yet. Set your first goal above!")
            else:
                for idx, g in enumerate(st.session_state.career_goals):
                    status_emoji = "⏳" if g["status"] == "Not Started" else ("🏃" if g["status"] == "In Progress" else "✅")
                    
                    goal_col1, goal_col2 = st.columns([4, 1])
                    with goal_col1:
                        st.markdown(f"**{status_emoji} {g['title']}** (Target: {g['date']})")
                    with goal_col2:
                        if g["status"] != "Completed":
                            if st.button("Mark Done", key=f"done_goal_{idx}"):
                                st.session_state.career_goals[idx]["status"] = "Completed"
                                st.success("Goal marked completed!")
                                st.rerun()
                        else:
                            st.write("Completed")
                            
        with l_col2:
            st.write("##### 🏆 Competition & Achievement Records")
            if not st.session_state.career_achievements:
                st.write("No competitions logged yet. Start tracking above!")
            else:
                for idx, c in enumerate(st.session_state.career_achievements):
                    st.markdown(f"""
                    <div class="premium-card" style="margin-bottom:0.5rem; padding:0.8rem;">
                        <b>{c["event"]} ({c["sport"]})</b>
                        <p style="margin: 0.1rem 0; font-size:0.85rem; color:#777;">📅 {c["date"]} | Level: {c["level"]}</p>
                        <p style="margin:0.2rem 0; font-weight:bold; color:#10b981;">Result: {c["result"]}</p>
                    </div>
                    """, unsafe_allow_html=True)

    # Sub-tab 7: Career FAQs
    with sub_career_tabs[6]:
        st.markdown("### ❓ Sports Career FAQs")
        st.write("Quick answers to frequently asked questions about sports careers, academies, eligibility, and scholarship programs.")
        
        faqs = [
            {
                "q": "What is the typical age limit to turn professional in sports?",
                "a": "For high-impact team sports like football or basketball, most professional players enter academies between ages 12-16 and sign professional contracts by ages 18-21. However, in sports like athletics, golf, or cricket, players can turn professional well into their mid-20s if their training metrics and state performances are outstanding."
            },
            {
                "q": "How do I get selected for sports scholarships?",
                "a": "Selection is merit-based. You must represent your school/college at district, state, or national-level competitions and win medals. Keep a certified folder of your participation certificates, newspaper clippings, and high-quality game videos. Apply directly to government portals (like SAI/Khelo India) or corporate sports foundations (JSW, Tata, Reliance) during their annual trials."
            },
            {
                "q": "What are the coaching certification levels, and where do I apply?",
                "a": "Coaching certifications are structured by national/international federations. For soccer, register via the AIFF/AFC licensing portal (starting with D-License). For cricket, contact your state cricket association for BCCI Level 1 coaching seminars. For general fitness coaching, look for certifications like NSCA-CSCS or NASM."
            },
            {
                "q": "Can I pursue sports management without playing sports at a high level?",
                "a": "Absolutely! Sports Management focuses on the business, legal, and operational aspects of sports. You do not need to be a professional athlete. Degrees like a Bachelor of Sports Management (BSM) or an MBA in Sports Management prepare you for careers in team operations, player agency, sports marketing, and event organization."
            },
            {
                "q": "How can I balance training 4 hours a day with school exams?",
                "a": "The key is strict scheduling and time blocking. Train early in the morning (5:00 AM - 7:00 AM) and late afternoon (4:30 PM - 6:30 PM). Use school study hall sessions efficiently and communicate with your school administration—many sports-focused academies or boards provide extensions or separate exam schedules for student-athletes."
            },
            {
                "q": "What are the opportunities in Sports Analytics, and what should I study?",
                "a": "Sports Analytics is an emerging field that involves evaluating performance data using statistical methods. You should focus on studying Data Science, Statistics, or Computer Science. Acquire technical skills in Python, SQL, Tableau, and data manipulation. Familiarize yourself with tracking databases like Opta, Wyscout, or CricViz."
            }
        ]
        
        faq_search = st.text_input("🔍 Search FAQs", placeholder="Type keywords (e.g. scholarship, age, coach, analytics)...")
        
        for idx, faq in enumerate(faqs):
            if not faq_search or faq_search.lower() in faq["q"].lower() or faq_search.lower() in faq["a"].lower():
                with st.expander(f"❓ {faq['q']}"):
                    st.write(faq["a"])

# ------------------------------------------------------------------
# Tab 4: Document QA (RAG)
# ------------------------------------------------------------------
with tab_rag:
    st.subheader("📂 Rulebook Semantic QA (RAG)")
    st.write("Upload any training manual, game guide, or official rulebook (PDF or Word document), and perform instant AI semantic searches through it.")

    if not HAS_RAG:
        st.error("⚠️ Document Q&A requires `sentence-transformers` and `faiss-cpu`. Please install them in the virtual environment to unlock this tab.")
    else:
        uploaded_file = st.file_uploader("Upload Document (.pdf or .docx)", type=["pdf", "docx"])
        
        if uploaded_file is not None:
            st.info("Parsing document contents...")
            
            # Extract text
            file_type = uploaded_file.name.split(".")[-1].lower()
            doc_text = ""
            
            try:
                if file_type == "pdf":
                    if HAS_PYPDF:
                        pdf_reader = pypdf.PdfReader(uploaded_file)
                        for page in pdf_reader.pages:
                            text_page = page.extract_text()
                            if text_page:
                                doc_text += text_page + "\n"
                    else:
                        st.error("Missing library 'pypdf' for PDF reading.")
                elif file_type == "docx":
                    if HAS_DOCX:
                        docx_file = docx.Document(uploaded_file)
                        doc_text = "\n".join([para.text for para in docx_file.paragraphs])
                    else:
                        st.error("Missing library 'python-docx' for DOCX reading.")
            except Exception as ex:
                st.error(f"Error parsing uploaded file: {ex}")

            if doc_text:
                st.success(f"Successfully extracted {len(doc_text)} characters.")
                
                # Check if already indexed in session state
                if st.session_state.get("rag_data") is None or st.session_state.rag_data.get("filename") != uploaded_file.name:
                    if st.button("⚡ Build Semantic Search Index"):
                        with st.spinner("Splitting document and generating vector embeddings (this might take a moment on first run)..."):
                            chunks = get_chunks(doc_text)
                            if embedding_model is not None:
                                try:
                                    # Encode & create FAISS index
                                    embeddings = embedding_model.encode(chunks)
                                    embeddings = np.array(embeddings).astype("float32")
                                    
                                    dim = embeddings.shape[1]
                                    index = faiss.IndexFlatL2(dim)
                                    index.add(embeddings)
                                    
                                    st.session_state.rag_data = {
                                        "filename": uploaded_file.name,
                                        "chunks": chunks,
                                        "index": index
                                    }
                                    st.success("🎉 Semantic search index built successfully!")
                                except Exception as index_err:
                                    st.error(f"Indexing failed: {index_err}")
                            else:
                                st.error("Embedding model could not be loaded.")
                
                # If indexed, show QA
                if st.session_state.get("rag_data") is not None:
                    st.markdown("---")
                    st.write("### 💬 Ask a question about the document:")
                    
                    rag_query = st.text_input("Enter your query:", placeholder="e.g. What are the rules for yellow cards? Or what is the drill frequency?")
                    k_val = st.slider("Context chunks to retrieve", min_value=1, max_value=5, value=3)
                    
                    if st.button("🔍 Search Document", key="rag_search_btn"):
                        if rag_query:
                            with st.spinner("Searching document database..."):
                                data = st.session_state.rag_data
                                index = data["index"]
                                chunks = data["chunks"]
                                
                                # Embed query
                                query_emb = embedding_model.encode([rag_query]).astype("float32")
                                _, indices = index.search(query_emb, k_val)
                                
                                # Retrieve matching chunks
                                matched_snippets = []
                                for idx in indices[0]:
                                    if 0 <= idx < len(chunks):
                                        matched_snippets.append(chunks[idx])
                                
                                # Display matches
                                st.write("### 📍 Top Matching Context Snippets")
                                for s_idx, snippet in enumerate(matched_snippets):
                                    st.markdown(f"""
                                    <div class="premium-card">
                                        <b style="color:#d946ef">Snippet {s_idx + 1}:</b>
                                        <p style="font-size:0.95rem; margin-top:0.4rem;">... {snippet} ...</p>
                                    </div>
                                    """, unsafe_allow_html=True)
                                
                                # Use LLM if configured
                                if active_api_key or use_mock:
                                    st.write("### 🤖 Coherent AI Answer")
                                    context_block = "\n---\n".join(matched_snippets)
                                    rag_prompt = (
                                        f"You are SportsAI, a document assistant.\n"
                                        f"Answer the user's question based strictly on the retrieved document snippets below.\n"
                                        f"If the information is not in the snippets, state that you cannot find it in the uploaded document.\n\n"
                                        f"Question: {rag_query}\n\n"
                                        f"Retrieved Document Snippets:\n{context_block}"
                                    )
                                    with st.spinner("Formulating answer..."):
                                        rag_reply = call_llm(rag_prompt, "You are a Document QA assistant.")
                                        st.markdown(rag_reply)
                        else:
                            st.warning("Please type a question before searching.")

# ------------------------------------------------------------------
# Tab 5: About SportsAI
# ------------------------------------------------------------------
with tab_about:
    st.subheader("ℹ️ About SportsAI Premium Assistant")
    
    st.markdown("""
    <div class="welcome-card" style="border-left: 5px solid #d946ef;">
        <h2 style="color: #d946ef; margin-top:0;">🏆 Elevating Athletic Excellence Through Intelligence</h2>
        <p style="font-size: 1.1rem; line-height: 1.7;">
            <b>SportsAI Pro</b> is an advanced assistant designed to provide elite-level coaching advice, tactical analysis, diet planning, rules breakdown, and document semantical search for athletes and trainers alike. By leveraging state-of-the-art Generative AI models from Google and Anthropic, SportsAI becomes your personal sports analyst, trainer, commentator, and fitness guide all in one dashboard.
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    col_ab1, col_ab2 = st.columns(2)
    with col_ab1:
        st.write("### 🛠️ Technology Architecture Stack")
        st.markdown("""
        * **Frontend Logic**: Built with **Streamlit** for interactive widgets, score updates, and real-time interval timers.
        * **AI Foundations**: Powered by **Google Gemini API** (e.g., `gemini-2.5-flash`) & **Anthropic Claude API** (e.g., `claude-3-5-sonnet-latest`).
        * **Semantic RAG Engine**: Utilizes **HuggingFace Sentence Transformers (`all-MiniLM-L6-v2`)** and **FAISS Vector DB** for offline-capable similarity search through rules.
        * **Export Framework**: Formats weekly workouts and diets into standard Microsoft Word docs using **python-docx**.
        """)
        
    with col_ab2:
        st.write("### 📈 Application Statistics")
        st.markdown(f"""
        <div style="display: grid; grid-template-columns: repeat(2, 1fr); gap: 1rem;">
            <div class="premium-card" style="text-align: center; margin-bottom: 0;">
                <b style="color: #ff1493;">Supported Sports</b>
                <h3 style="margin: 0.3rem 0; color: #ff1493;">20+</h3>
            </div>
            <div class="premium-card" style="text-align: center; margin-bottom: 0;">
                <b style="color: #06b6d4;">RAG Parsing</b>
                <h3 style="margin: 0.3rem 0; color: #06b6d4;">PDF & DOCX</h3>
            </div>
            <div class="premium-card" style="text-align: center; margin-bottom: 0;">
                <b style="color: #d946ef;">Active Theme</b>
                <h3 style="margin: 0.3rem 0; color: #d946ef;">{st.session_state.current_theme}</h3>
            </div>
            <div class="premium-card" style="text-align: center; margin-bottom: 0;">
                <b style="color: #10b981;">AI Engines</b>
                <h3 style="margin: 0.3rem 0; color: #10b981;">Gemini & Claude</h3>
            </div>
        </div>
        """, unsafe_allow_html=True)
        
    st.markdown("---")
    st.write("### 🧭 Quick Usage Guide")
    cols_g = st.columns(3)
    with cols_g[0]:
        st.markdown("""
        **1. Configure API Key**
        Add your Google Gemini API Key in the left Control Panel. It will auto-save to `.env` instantly without errors.
        """)
    with cols_g[1]:
        st.markdown("""
        **2. Upload Rulebooks**
        Navigate to the *Rulebook Semantic QA* tab, drag and drop any sports manual PDF, and query it directly.
        """)
    with cols_g[2]:
        st.markdown("""
        **3. Export Workouts**
        Fill in your sports interests in the *Workout & Diet Planner*, hit Build, and download your schedule as a `.docx` file.
        """)